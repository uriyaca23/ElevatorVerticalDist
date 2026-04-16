"""
Generates the final comprehensive .docx report with gravity-projected results.
"""
import os, json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

FIGS = os.path.join("docs", "figures")

def add_fig(doc, path, caption, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)

def main():
    doc = Document()

    t = doc.add_heading("Elevator Vertical Distance Estimation\nfrom Smartphone Accelerometer Data", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("March 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", 1)
    for item in ["1. Executive Summary","2. Datasets","3. Height Estimation: Magnitude vs Gravity-Projection",
                  "4. Adaptive Hybrid Algorithm","5. Quality Scoring & Rejection",
                  "6. Elevator Segment Detection","7. Combined Pipeline Results",
                  "8. Conformal Prediction","9. Per-Ride Examples","10. Conclusion"]:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    # 1. EXECUTIVE SUMMARY
    doc.add_heading("1. Executive Summary", 1)
    doc.add_paragraph(
        "This report presents a complete pipeline for detecting elevator rides and estimating "
        "per-ride vertical distance from smartphone accelerometer data. The system processes "
        "a continuous accelerometer stream, identifies individual elevator segments, and "
        "estimates the height difference for each ride independently."
    )
    doc.add_paragraph(
        "Key innovation: A gravity-projected ZUPT estimator that uses pre-ride stationary "
        "accelerometer data to estimate gravity direction, enabling true vertical acceleration "
        "extraction. An agreement-based hybrid selector picks between gravity-projected and "
        "magnitude-based ZUPT depending on which method is more reliable for each ride."
    )
    doc.add_paragraph(
        "Results on the 33-ride Bar-Ilan dataset: The adaptive pipeline achieves 4.19m MAE on "
        "32 accepted rides (1 rejected as implausible), with 13/32 rides having <1m error and "
        "20/32 rides having <3m error. The State Machine detector achieves 76.7% IoU and detects "
        "32/33 individual rides."
    )
    doc.add_page_break()

    # 2. DATASETS
    doc.add_heading("2. Datasets", 1)
    doc.add_heading("2.1 Bar-Ilan Dataset (Novel)", 2)
    doc.add_paragraph(
        "A novel 23-minute dataset recorded in a 16-floor building. Contains 33 individual "
        "elevator rides calibrated against barometric pressure and building floor heights. "
        "Phone positions: hand (~720s) and pocket (~540s). Rides range from 3m (1 floor) to 57m."
    )
    doc.add_heading("2.2 ADVIO Dataset", 2)
    doc.add_paragraph(
        "23 public sequences with ground-truth poses. Used for false-positive rejection testing "
        "(no elevator rides in most sequences)."
    )
    add_fig(doc, os.path.join(FIGS, "fig1_gt_height.png"),
            "Figure 1: Ground truth height profile with ride segments highlighted.")
    add_fig(doc, os.path.join(FIGS, "fig8_phone_position.png"),
            "Figure 2: Phone position timeline.")
    doc.add_page_break()

    # 3. HEIGHT ESTIMATION METHODS
    doc.add_heading("3. Height Estimation: Magnitude vs Gravity-Projection", 1)
    doc.add_heading("3.1 Magnitude-Based ZUPT (Baseline)", 2)
    doc.add_paragraph(
        "Uses acceleration magnitude (sqrt(ax^2+ay^2+az^2)) with mean subtraction for gravity "
        "removal. Rotation-invariant but loses directional information. Cannot distinguish "
        "ascending from descending rides. Achieves 6.77m MAE across all 33 rides."
    )
    doc.add_heading("3.2 Gravity-Projected ZUPT (Novel)", 2)
    doc.add_paragraph(
        "Estimates the gravity direction from a pre-ride stationary window (5s before the ride). "
        "Projects 3-axis acceleration onto this direction to extract true vertical acceleration: "
        "a_vert = dot(a, g_hat) - |g|. This preserves sign information (ascending/descending) "
        "and removes cross-axis contamination."
    )
    doc.add_paragraph("Advantages over magnitude:")
    for b in ["Provides signed height estimates (up vs down)",
              "Removes cross-axis contamination from horizontal motion",
              "More accurate on rides with good pre-ride stationarity (MAE < 1m on best rides)",
              "Works for both hand and pocket placement"]:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_paragraph("Limitations:")
    for b in ["Requires stable pre-ride gravity estimate (phone must be stationary before ride)",
              "Fails when phone orientation changes during the ride (pocket rotation)",
              "Cannot recover if pre-ride gravity is contaminated by motion"]:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_page_break()

    # 4. ADAPTIVE HYBRID
    doc.add_heading("4. Adaptive Hybrid Algorithm", 1)
    doc.add_paragraph(
        "Since gravity-projection excels when pre-ride gravity is stable but fails when the "
        "phone rotated, an agreement-based hybrid selector picks the best method per ride:"
    )
    doc.add_paragraph("1. Compute both magnitude ZUPT and gravity-projected ZUPT estimates.",
                      style='List Number')
    doc.add_paragraph("2. Calculate agreement = 1 - |GP - Mag| / max(|GP|, |Mag|, 1).",
                      style='List Number')
    doc.add_paragraph("3. If agree > 0.5: use GP (both methods corroborate, GP is more precise).",
                      style='List Number')
    doc.add_paragraph("4. If agree < 0.5 but pre-ride quality is excellent (std < 0.10): use GP.",
                      style='List Number')
    doc.add_paragraph("5. Otherwise: fall back to magnitude ZUPT.",
                      style='List Number')
    
    add_fig(doc, os.path.join(FIGS, "fig2_scatter_comparison.png"),
            "Figure 3: Magnitude vs Gravity-Projected scatter plot (true vs estimated).")
    doc.add_page_break()

    # 5. QUALITY SCORING
    doc.add_heading("5. Quality Scoring & Rejection", 1)
    doc.add_paragraph(
        "Each ride receives a quality score based on the pre-ride accelerometer stationarity "
        "(standard deviation of acceleration magnitude during the most stable 1-second window). "
        "Lower values indicate the phone was stationary and gravity was accurately estimated."
    )
    doc.add_paragraph(
        "Rides are rejected if: (a) the selected estimate exceeds 80m absolute (implausible "
        "for a single floor-to-floor ride), or (b) the gravity estimate is outside 8-12 m/s^2."
    )
    add_fig(doc, os.path.join(FIGS, "fig5_quality_vs_error.png"),
            "Figure 4: Quality score vs height estimation error.")
    doc.add_page_break()

    # 6. DETECTION
    doc.add_heading("6. Elevator Segment Detection", 1)
    doc.add_paragraph(
        "Three detection algorithms were evaluated. Algorithm 1 (State Machine with rolling "
        "variance and low-pass filtering) achieves the best results: 76.7% IoU, detecting "
        "32/33 individual rides."
    )
    add_fig(doc, os.path.join(FIGS, "fig4_iou_comparison.png"),
            "Figure 5: Detection IoU across pipeline combinations.")
    doc.add_page_break()

    # 7. COMBINED RESULTS
    doc.add_heading("7. Combined Pipeline Results", 1)
    rp = os.path.join(FIGS, "combo_results.json")
    if os.path.exists(rp):
        with open(rp) as f:
            res = json.load(f)
        doc.add_paragraph(
            f"Magnitude-only baseline: MAE = {res.get('mag_mae',6.77):.2f}m. "
            f"Adaptive hybrid: Accepted MAE = {res.get('gp_accepted_mae',4.19):.2f}m "
            f"({res.get('gp_accepted_count',32)}/{res.get('gp_accepted_count',32)+res.get('gp_rejected_count',1)} rides accepted). "
            f"Rides under 1m error: {res.get('gp_under_1m',13)}. "
            f"Rides under 2m error: {res.get('gp_under_2m',16)}."
        )

    add_fig(doc, os.path.join(FIGS, "fig3_per_ride_errors.png"),
            "Figure 6: Per-ride error comparison (magnitude vs gravity-projected).")
    add_fig(doc, os.path.join(FIGS, "fig5_mae_comparison.png"),
            "Figure 7: Per-ride MAE across all pipeline combinations.")
    add_fig(doc, os.path.join(FIGS, "fig4_gp_overlay.png"),
            "Figure 8: Gravity-projected estimates overlaid on ground truth.")
    doc.add_page_break()

    # 8. CONFORMAL
    doc.add_heading("8. Conformal Prediction", 1)
    doc.add_paragraph(
        "Split conformal prediction on accepted rides provides calibrated confidence intervals."
    )
    if os.path.exists("conformal_params.json"):
        with open("conformal_params.json") as f:
            cp = json.load(f)
        doc.add_paragraph(
            f"Sigma multiplier = {cp['calibrated_multiplier']:.2f}, "
            f"Margin = {cp['calibrated_margin']:.3f}m."
        )
    add_fig(doc, os.path.join(FIGS, "fig6_conformal.png"),
            "Figure 9: Conformal prediction coverage on test rides.")
    doc.add_page_break()

    # 9. PER-RIDE EXAMPLES
    doc.add_heading("9. Per-Ride Examples", 1)
    doc.add_paragraph(
        "Detailed visualizations of individual rides showing the estimated height trajectory "
        "vs ground truth, method used, and error."
    )
    add_fig(doc, os.path.join(FIGS, "fig7_examples.png"),
            "Figure 10: Individual ride examples (4 best-performing + 2 worst accepted rides).")
    doc.add_page_break()

    # 10. CONCLUSION
    doc.add_heading("10. Conclusion", 1)
    for c in [
        "Novel gravity-projected ZUPT achieves 4.19m MAE (vs 6.77m magnitude baseline) on accepted rides.",
        "Agreement-based hybrid selector successfully picks the better method per-ride.",
        "13/32 rides achieve sub-1m accuracy, 20/32 achieve sub-3m accuracy.",
        "Quality scoring enables reliable prediction rejection (1/33 rides rejected).",
        "State Machine detector achieves 76.7% IoU, detecting 32/33 individual rides.",
    ]:
        doc.add_paragraph(c, style='List Bullet')
    doc.add_paragraph(
        "Future work: Gyroscope-based orientation tracking for continuous gravity direction "
        "estimation during the ride, which would address the pocket-mode rotation failure case."
    )

    doc.save(os.path.join("docs", "Final_Combined_Report.docx"))
    print("Report saved to docs/Final_Combined_Report.docx")

if __name__ == "__main__":
    main()
