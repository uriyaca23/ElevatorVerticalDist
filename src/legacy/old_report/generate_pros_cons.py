import os
from docx import Document
from docx.shared import Pt

def create_doc():
    doc = Document()
    doc.add_heading('Algorithms for Estimating Elevator Vertical Distance', 0)

    doc.add_paragraph('This document outlines the three algorithms chosen to estimate the vertical distance travelled by an elevator using only accelerometer data. It discusses their theoretical background, pros, and cons.')

    # Algorithm 1
    doc.add_heading('1. Naive Double Integration', level=1)
    doc.add_paragraph('Theory: The simplest approach. It assumes the initial velocity is zero. It subtracts a static estimate of gravity from the Z-axis acceleration, then uses the trapezium rule to numerically integrate acceleration into velocity, and then velocity into position.')
    doc.add_heading('Pros:', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Extremely simple to implement and computationally inexpensive.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Requires no complex state tuning or thresholding parameters.')
    doc.add_heading('Cons:', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Highly susceptible to sensor noise and gravity bias estimation errors.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Quadratic drift: Any constant bias in acceleration leads to a linear drift in velocity and a quadratic drift in position, rendering it largely inaccurate for long rides.')

    # Algorithm 2
    doc.add_heading('2. Double Integration with Zero Velocity Update (ZUPT)', level=1)
    doc.add_paragraph('Theory: Elevators have distinct stationary periods. By calculating the rolling variance of the acceleration signal, we can detect when the elevator is at rest. During these periods, we know the true velocity is exactly 0 m/s. The ZUPT algorithm numerically integrates the acceleration but forces the velocity to 0 during rest. Furthermore, the accumulated drift error at the end of a motion segment is linearly distributed backwards across the segment to correct the velocity profile before integrating it to distance.')
    doc.add_heading('Pros:', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Effectively eliminates long-term drift by regularly resetting the error bounds.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Drift distribution greatly increases the accuracy of the final position estimate.')
    doc.add_heading('Cons:', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Requires careful tuning of the variance threshold to accurately detect stationary periods without false positives (e.g., mistaking a constant-velocity phase for being at rest).')

    # Algorithm 3
    doc.add_heading('3. 1D Kalman Filter with Kinematic State Constraints', level=1)
    doc.add_paragraph('Theory: A Kalman filter estimates the state of a linear dynamic system. Our state vector includes Position, Velocity, and Acceleration Bias. The propagation model uses the raw accelerometer readings as inputs. The observation model relies on pseudo-measurements: when the variance of the acceleration indicates the elevator is stationary, we feed a velocity measurement of 0 m/s with very low uncertainty into the filter. This allows the filter to continuously estimate and correct the dynamic sensor bias.')
    doc.add_heading('Pros:', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Mathematically robust method for dynamically estimating and removing sensor bias on-the-fly.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Provides not just estimates, but also uncertainty bounds (covariance) for the distance.')
    doc.add_heading('Cons:', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Computationally heavier than direct integration.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Requires careful tuning of the Process Noise (Q) and Measurement Noise (R) covariance matrices to achieve optimal performance.')

    os.makedirs(r"c:\Users\uriya\PycharmProjects\ElevatorVerticalDist\docs", exist_ok=True)
    out_path = r"c:\Users\uriya\PycharmProjects\ElevatorVerticalDist\docs\algorithms_pros_cons.docx"
    doc.save(out_path)
    print(f"Document saved to {out_path}")

if __name__ == "__main__":
    create_doc()
