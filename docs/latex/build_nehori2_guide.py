"""Build the Hebrew Word operating guide "מדריך הפעלה של נהורי 2".

A concise, image-based operator manual for the Boutique Pipeline Streamlit app.
Reuses the existing screenshots under docs/latex/figures/boutique/ plus two generated
assets under docs/latex/figures/nehori2/ (pipeline flow diagram + extracted S-curve figure).

Run with:  venv/bin/python docs/latex/build_nehori2_guide.py
Output:    docs/latex/nehori2_guide_hebrew.docx
"""
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
FIG = os.path.join(ROOT, "docs", "latex", "figures")
BQ = os.path.join(FIG, "boutique")
NH = os.path.join(FIG, "nehori2")
OUT = os.path.join(ROOT, "docs", "latex", "nehori2_guide_hebrew.docx")

HEB_FONT = "Arial"
NAVY = RGBColor(0x15, 0x32, 0x4F)
RED = RGBColor(0xB0, 0x3A, 0x2E)
GRAY = RGBColor(0x6B, 0x6B, 0x6B)


# ---------------------------------------------------------------- RTL helpers
def _set_rtl_paragraph(p):
    pPr = p._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")


def _style_run(run, size=12, bold=False, color=None, font=HEB_FONT):
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.name = font
    rPr = run._r.get_or_add_rPr()
    # complex-script + ascii fonts
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    # mark run as RTL
    rtl = rPr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)
    rtl.set(qn("w:val"), "1")
    # complex-script size
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(int(size * 2)))


def para(doc, segments, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT,
         space_before=2, space_after=8, color=NAVY):
    """segments: str or list of (text, bold) tuples."""
    if isinstance(segments, str):
        segments = [(segments, False)]
    p = doc.add_paragraph()
    p.alignment = align
    _set_rtl_paragraph(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    for text, bold in segments:
        run = p.add_run(text)
        _style_run(run, size=size, bold=bold, color=color)
    return p


def heading(doc, text, size=16, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl_paragraph(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _style_run(run, size=size, bold=True, color=color)
    return p


def numbered(doc, n, segments, size=12):
    if isinstance(segments, str):
        segments = [(segments, False)]
    segments = [(f"{n}.  ", True)] + list(segments)
    return para(doc, segments, size=size, space_after=6)


def image(doc, path, width_in):
    assert os.path.exists(path), f"missing image: {path}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl_paragraph(p)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    _style_run(run, size=10, bold=False, color=GRAY)
    run.italic = True
    return p


# ---------------------------------------------------------------- document
doc = Document()
sec = doc.sections[0]
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)
sec.top_margin = Inches(0.9)
sec.bottom_margin = Inches(0.9)
USABLE = sec.page_width.inches - 1.8  # ~6.7 in on letter

# base style font
normal = doc.styles["Normal"]
normal.font.name = HEB_FONT
normal.font.size = Pt(12)

# ===================== PAGE 1 — TITLE =====================
for _ in range(6):
    doc.add_paragraph()
para(doc, [("מדריך הפעלה של נהורי 2", True)], size=34,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, color=NAVY)
para(doc, [("הערכת מרחק אנכי במעלית מרשומת חיישני סמארטפון", False)], size=15,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, color=GRAY)
for _ in range(11):
    doc.add_paragraph()
para(doc, [("נוצר ע״י אייל יקיר ואוריה כהן", True)], size=14,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, color=NAVY)
para(doc, [("אחראי לאשר תוצרים חריגים — אייל יקיר, אוריה כהן, גל מחלין", False)],
     size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
doc.add_page_break()

# ===================== PAGE 2 — OVERVIEW DIAGRAM =====================
heading(doc, "סקירת התהליך: ארבעת שלבי האלגוריתם", size=18)
para(doc,
     "הפייפליין מורכב מארבעה שלבים שרצים בזה אחר זה. כל שלב מעביר את התוצר שלו לשלב הבא, "
     "עד שמתקבל דו״ח סופי לכל נסיעת מעלית שזוהתה.")
image(doc, os.path.join(NH, "pipeline_flow.png"), min(USABLE, 6.6))
caption(doc, "תרשים 1: זרימת ארבעת השלבים. החיצים האדומים מסמנים את השלבים שאינם אוטומטיים לחלוטין.")
para(doc, [
    ("שני השלבים הראשון והאחרון — ", False),
    ("שליפת נתונים", True),
    (" ו", False),
    ("הפקת הדו״ח", True),
    (" — רצים אוטומטית ואינם דורשים שיקול דעת. לעומתם, ", False),
    ("סגמנטציה", True),
    (" ו", False),
    ("חיזוי", True),
    (" הם שלבים חצי-אוטומטיים: האלגוריתם מציע פתרון, אך על המשתמש לעבור עליו, לאמת אותו, "
     "ולתקן אותו במידת הצורך לפני המעבר הלאה. עיקר תשומת הלב של המפעיל מתמקדת בשני שלבים אלו.", False),
])
doc.add_page_break()

# ===================== DATA LOADING =====================
heading(doc, "שלב 1 — שליפת נתונים", size=18)
para(doc,
     "בשלב זה בוחרים את מקור הנתונים. למשתמשים פנימיים, לוחצים על "
     "‏Use phone DB‏ כדי לשלוף הקלטה ישירות ממאגר הניסויים.")
image(doc, os.path.join(BQ, "ui_step2_picker.png"), min(USABLE, 6.6))
caption(doc, "תרשים 2: דף בחירת מקור הנתונים. בוחרים ב-Use phone DB לשליפה ממאגר.")
para(doc, [("לאחר הבחירה ב-Use phone DB ממלאים שלושה שדות:", True)])
numbered(doc, 1, [("מזהה הטלפון ", True), ("(phone ID)", False),
                  (" — סוג ומזהה הטלפון שביצע את ההקלטה.", False)])
numbered(doc, 2, [("תאריך ושעת התחלה", True),
                  (" — תחילת חלון הזמן שממנו לשלוף את הנתונים.", False)])
numbered(doc, 3, [("תאריך ושעת סיום", True),
                  (" — סוף חלון הזמן.", False)])
para(doc, [("בסיום לוחצים על ", False), ("Next", True), (" כדי לשלוף את הנתונים ולעבור לשלב הסגמנטציה.", False)])
para(doc, [
    ("חשוב: ", True),
    ("יש לבחור חלונות זמן של ", False),
    ("עד שעה אחת לכל היותר", True),
    (". שליפת הנתונים מהמאגר אורכת זמן רב, וחלון ארוך מדי יאריך מאוד את ההמתנה. "
     "אם דרוש טווח גדול יותר — מומלץ לפצל אותו למספר שליפות נפרדות.", False),
])
doc.add_page_break()

# ===================== SEGMENTATION — THEORY =====================
heading(doc, "שלב 2 — סגמנטציה: מה אנחנו מחפשים", size=18)
para(doc,
     "מטרת הסגמנטציה היא לאתר בתוך ההקלטה הרציפה את קטעי הזמן שבהם התרחשה נסיעת מעלית, "
     "ולתייג כל נסיעה כעלייה או ירידה.")
para(doc, [
    ("הזיהוי מתבסס על ", False),
    ("חתימת התאוצה האופיינית של מעלית", True),
    (". מעלית שמאיצה ואז מאטה יוצרת ", False),
    ("שתי אונות (פולסים)", True),
    (" — אונת האצה ראשונה ואונת האטה שנייה — שיחד נראות כעקומת ", False),
    ("S", False),
    (" (טרפז כפול). זו התבנית שהגלאי מחפש לאורך כל האות.", False),
])
para(doc,
     "צורת העקומה תלויה באורך הנסיעה (מספר הקומות). ככל שהנסיעה ארוכה יותר, כך המעלית "
     "מספיקה למצות יותר ממגבלות התנועה שלה. נבחין בשלושה משטרים:")
image(doc, os.path.join(NH, "scurve_regimes.png"), min(USABLE, 6.2))
caption(doc, "תרשים 3: שלושת המשטרים של עקומת ה-S (מתוך המאמר). נסיעה ארוכה — 7 קטעים, "
             "עם רמת שיוט שטוחה בין שתי האונות; נסיעה בינונית — 5 קטעים, שתי האונות נוגעות; "
             "נסיעה קצרה — 3 קטעים, שתי אונות משולשיות.")
para(doc,
     "המשותף לשלושת המשטרים: אונה חיובית אחת ואחריה אונה שלילית (או להפך), כשביניהן האות "
     "חוזר לאפס. זו בדיוק הצורה שיש לחפש כשבודקים אם מקטע שזוהה הוא אכן נסיעת מעלית אמיתית.")
doc.add_page_break()

# ===================== SEGMENTATION — OUTPUT & TASKS =====================
heading(doc, "שלב 2 — תוצאת הסגמנטציה ותפקיד המשתמש", size=18)
para(doc,
     "הגלאי רץ אוטומטית ומסמן את הנסיעות שזיהה. כל נסיעה מופיעה כרצועה צבעונית מעל האות: "
     "כחול = עלייה, ורוד = ירידה, והמקטע שנבחר כרגע מסומן בכתום.")
image(doc, os.path.join(BQ, "panel_clean_timeline.png"), min(USABLE, 6.6))
caption(doc, "תרשים 4: תצוגת הסגמנטציה באתר (פאנל ‏Signal + segments‏) — האות עם רצועות הנסיעה הצבעוניות.")
para(doc, [("הגלאי אינו מושלם. תפקיד המשתמש בשלב זה הוא לעבור על הפלט ולבצע שלוש בדיקות:", True)])

numbered(doc, 1, [
    ("בדיקת כל מקטע שזוהה. ", True),
    ("בוחרים כל נסיעה ובודקים בגרף הטרפז המותאם אם היא אכן בעלת צורת עקומת ה-S "
     "(שתי אונות ברורות). אם המקטע אינו בצורה הזו — מדובר בזיהוי שגוי, ויש ", False),
    ("למחוק", True),
    (" אותו.", False),
])
image(doc, os.path.join(BQ, "panel_clean_trapezoid.png"), min(USABLE, 5.6))
caption(doc, "התאמת טרפז תקינה (משאירים את המקטע): שתי האונות יושבות על האות.")
image(doc, os.path.join(BQ, "panel_fp_trapezoid.png"), min(USABLE, 5.6))
caption(doc, "התאמה שגויה (מוחקים את המקטע): אין צורת שתי-אונות אמיתית — קוץ/רעש בלבד.")

numbered(doc, 2, [
    ("איתור נסיעות שפוספסו. ", True),
    ("סורקים את האות בעיניים ומחפשים צורות שתי-אונות ברורות שאין מעליהן רצועה צבעונית — "
     "אלו נסיעות שהגלאי החמיץ. כל נסיעה כזו יש ", False),
    ("להוסיף", True),
    (" ידנית.", False),
])
image(doc, os.path.join(BQ, "panel_damped_timeline.png"), min(USABLE, 6.6))
caption(doc, "דוגמה להקלטה שבה זוהו מעט מקטעים אך בפועל יש נסיעות נוספות שלא סומנו — יש להוסיפן.")

numbered(doc, 3, [
    ("תיקון מקטעים שאוחדו. ", True),
    ("לעיתים הגלאי מאחד שתי נסיעות נפרדות למקטע אחד. למשל ברצף של אונה-עלייה, אונה-ירידה, "
     "אונה-עלייה, אונה-ירידה — הגלאי עלול לתפוס בטעות את שתי האונות האמצעיות (ירידה ואז עלייה) "
     "כמקטע יחיד, במקום לזהות שתי נסיעות נפרדות. במקרה כזה ", False),
    ("מוחקים את המקטע האמצעי השגוי ויוצרים שני מקטעים חדשים", True),
    (" במקומו.", False),
])
image(doc, os.path.join(BQ, "panels", "haari_step3_seg20_full__trap.png"), min(USABLE, 6.4))
caption(doc, "מקטע יחיד ארוך מדי (537–558 ש׳) שמכסה שני אירועים: אונה 1 (אדום) תופסת המראה ב-538 ש׳, "
             "אונה 2 (סגול) נחיתה ב-557 ש׳, עם ~17 שניות שקטות באמצע — יש לפצל לשתי נסיעות.")

para(doc, [
    ("פקדי העריכה ", False),
    ("(בסרגל הצד): ", True),
    ("Add segment +", True),
    (" להוספת מקטע, ", False),
    ("Edit selected", True),
    (" לעריכת זמני ההתחלה/סיום והסוג, ", False),
    ("Delete", True),
    (" (אייקון הפח) למחיקה, ו-", False),
    ("Reset", True),
    (" לאיפוס כל העריכות חזרה לפלט המקורי של הגלאי.", False),
])
para(doc, [
    ("כשהמקטעים תקינים, לוחצים על ", False),
    ("Predict", True),
    (" כדי לעבור לשלב החיזוי.", False),
])
doc.add_page_break()

# ===================== PREDICTION =====================
heading(doc, "שלב 3 — חיזוי", size=18)
para(doc,
     "בשלב החיזוי מחושב לכל מקטע מרחק אנכי משוער (Δh) במטרים. החישוב רץ אוטומטית, "
     "ותפקיד המשתמש מתמקד בדבר אחד: לבדוק את גרף הטרפז המותאם של כל מקטע.")
image(doc, os.path.join(BQ, "panel_clean_trapezoid.png"), min(USABLE, 5.8))
caption(doc, "תרשים 5: גרף הטרפז המותאם. בודקים שהתבנית יושבת היטב על האות בשתי האונות.")
para(doc, [
    ("עוברים על כל מקטע ובודקים אם יש ", False),
    ("רעש או בעיה באחת האונות", True),
    (" — אם תבנית הטרפז אינה יושבת היטב על האות. אם זה המצב, מתקנים ידנית את פרמטרי הטרפז "
     "בכרטיס הפרמטרים: ", False),
    ("f", True),
    (" (שבר הרמה השטוחה), ", False),
    ("W", True),
    (" (חצי-הרוחב), ו-", False),
    ("A", True),
    (" (המשרעת) — עד שהתבנית מתאימה לאות.", False),
])
image(doc, os.path.join(BQ, "crops", "clean_step3_overview__params.png"), min(USABLE, 6.4))
caption(doc, "תרשים 6: כרטיס פרמטרי הטרפז (Fitted trapezoid parameters) — כאן עורכים את f, W ו-A "
             "לכל אונה (LOBE 1 / LOBE 2).")
para(doc, [
    ("לאחר התיקון לוחצים על ", False),
    ("Re-predict", True),
    (" כדי לחשב מחדש את החיזוי עם הפרמטרים המעודכנים, ואז ", False),
    ("Next", True),
    (".", False),
])
para(doc, [
    ("כשמסיימים לעבור על כל המקטעים, לוחצים על ", False),
    ("Generate report", True),
    (" — וזהו. מתקבל דו״ח PDF סופי עם המרחק האנכי המשוער לכל נסיעה.", False),
])
image(doc, os.path.join(BQ, "clean_step5_report.png"), min(USABLE, 6.2))
caption(doc, "תרשים 7: תצוגה מקדימה של דו״ח ה-PDF הסופי.")

doc.save(OUT)
print("wrote", OUT)
print("paragraphs:", len(doc.paragraphs))
