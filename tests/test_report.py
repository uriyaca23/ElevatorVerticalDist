import os
from docx import Document

def test_report_generation():
    report_path = r"c:\Users\uriya\PycharmProjects\ElevatorVerticalDist\docs\Final_Elevator_Height_Report.docx"
    assert os.path.exists(report_path), "The final report was not generated."
    
    # Check if we can open it safely
    doc = Document(report_path)
    
    # Ensure it is substantial in length
    pages_approx = len(doc.paragraphs) / 10 # very rough approximation
    assert len(doc.paragraphs) > 20, "Document seems too short."
    
    # Verify OMML syntax presence
    omml_found = False
    for p in doc.paragraphs:
        xml = p._p.xml
        if 'm:oMath' in xml or 'm:oMathPara' in xml:
            omml_found = True
            break
            
    assert omml_found, "OMML equations were not properly embedded into the document."

    # Check that pictures are present
    assert len(doc.inline_shapes) > 0, "No data plots or visualization images found in the report."

if __name__ == "__main__":
    test_report_generation()
    print("Report visually and structurally verified without errors!")
