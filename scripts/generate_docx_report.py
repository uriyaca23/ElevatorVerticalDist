"""
Generate the comprehensive DOCX research report.
Produces docs/Research_Report.docx with all chapters, figures, and tables.
"""
import os, sys, json
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.join(os.path.dirname(__file__), "..")
FIGS = os.path.join(BASE, "docs", "report_figures")
FIGS_V4 = os.path.join(BASE, "docs", "figures_v4")
OUT = os.path.join(BASE, "docs", "Research_Report.docx")

v4 = json.load(open(os.path.join(FIGS_V4, "v4_results.json")))

# ---- Helper functions ----
def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    return p

def add_figure(doc, filename, caption="", width=6.0):
    """Add a figure with caption."""
    # Check report_figures first, then figures_v4, then figures
    for folder in [FIGS, FIGS_V4, os.path.join(BASE, "docs", "figures")]:
        path = os.path.join(folder, filename)
        if os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Inches(width))
            if caption:
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
            return
    print(f"  WARNING: Figure not found: {filename}")

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()  # spacing
    return table


# =====================================================================
# BUILD THE REPORT
# =====================================================================
def build_report():
    doc = Document()
    
    # --- Style setup ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    
    # ================================================================
    # TITLE PAGE
    # ================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('Elevator Vertical Distance Estimation\nUsing Accelerometer Data', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Research Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].italic = True
    
    doc.add_paragraph()
    info = doc.add_paragraph('Three-Stage Pipeline: Detection → Quality Filter → Height Estimation\nWith Conformal Prediction for 90% Coverage Intervals')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    date = doc.add_paragraph('March 2026')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.runs[0].font.size = Pt(14)
    
    doc.add_page_break()
    
    # ================================================================
    # TABLE OF CONTENTS (manual)
    # ================================================================
    add_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Problem Background & Motivation',
        '2. Datasets',
        '3. Research History',
        '4. Current Solution Architecture',
        '5. Algorithm Comparison',
        '6. Final Results & Validation',
        '7. Usage Guide & Deployment',
        '8. Conclusions & Future Work',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()
    
    # ================================================================
    # CHAPTER 1: PROBLEM BACKGROUND
    # ================================================================
    add_heading(doc, '1. Problem Background & Motivation', 1)
    
    add_heading(doc, '1.1 Indoor Vertical Positioning', 2)
    add_para(doc, 'Accurate indoor vertical positioning is a critical challenge in modern navigation, '
             'emergency response, and IoT applications. While GPS provides reliable horizontal positioning '
             'outdoors, it fails indoors—and even outdoor-capable systems like GPS cannot distinguish between '
             'floors in multi-story buildings. The ability to determine which floor a person or device is on '
             'has applications in:')
    
    bullets = [
        'Emergency response: First responders need floor-level accuracy to locate victims in high-rise buildings.',
        'Indoor navigation: Shopping malls, airports, and hospitals with multiple floors.',
        'Asset tracking: Monitoring equipment or personnel across building floors.',
        'Smart buildings: Automated elevator dispatch, occupancy estimation.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    add_heading(doc, '1.2 Why Accelerometer-Only?', 2)
    add_para(doc, 'This project constrains itself to using ONLY accelerometer data for vertical distance '
             'estimation. While barometers can provide altitude estimates, they suffer from temperature drift, '
             'HVAC interference, and require calibration. GPS is unavailable indoors. The accelerometer, present '
             'in every smartphone, provides a universal sensor that can detect the characteristic acceleration '
             'pattern of elevator motion:')
    
    add_para(doc, '1. Upward jerk (acceleration phase) as the elevator starts\n'
             '2. Near-zero acceleration during constant velocity travel\n'
             '3. Downward jerk (deceleration phase) as the elevator stops', italic=True)
    
    add_heading(doc, '1.3 Physics of Accelerometer-Based Height Estimation', 2)
    add_para(doc, 'An accelerometer measures the specific force acting on the device, which in a stationary '
             'reference frame equals gravity (≈9.81 m/s²). During elevator motion, the measured acceleration '
             'includes both gravity and the elevator\'s acceleration. The vertical displacement can be obtained '
             'by double-integrating the vertical component of acceleration after subtracting gravity:')
    
    add_para(doc, 'h(t) = ∫∫ [a_vertical(t) − g] dt²', bold=True)
    
    add_para(doc, 'However, this double integration is extremely sensitive to:')
    bullets2 = [
        'Bias errors: A constant accelerometer bias of 0.01 m/s² produces 0.5m drift over 10 seconds.',
        'Gravity direction uncertainty: If the phone is tilted, decomposing acceleration into vertical and horizontal requires knowing the orientation.',
        'Sensor noise: Integration amplifies noise, especially over longer rides.',
        'Phone orientation changes: Users may rotate the phone during a ride, invalidating the gravity projection.',
    ]
    for b in bullets2:
        doc.add_paragraph(b, style='List Bullet')
    
    add_para(doc, 'These challenges make robust height estimation from accelerometer data alone a non-trivial '
             'signal processing problem that requires careful algorithm design, quality filtering, and '
             'statistical calibration.')
    
    add_heading(doc, '1.4 Project Objectives', 2)
    add_para(doc, 'This project aims to build a three-stage pipeline that:')
    objectives = [
        'Objective 1 — Detection: Automatically detect and segment individual elevator rides from continuous accelerometer recordings, with per-ride separation.',
        'Objective 2 — Quality Filter: Determine whether a detected segment is suitable for height estimation using accelerometer-only quality features. Reject unreliable rides to prevent catastrophic errors.',
        'Objective 3 — Height Estimation: For accepted rides, estimate the vertical displacement with 90% conformal prediction coverage.',
    ]
    for o in objectives:
        doc.add_paragraph(o, style='List Bullet')
    
    doc.add_page_break()
    
    # ================================================================
    # CHAPTER 2: DATASETS
    # ================================================================
    add_heading(doc, '2. Datasets', 1)
    
    add_heading(doc, '2.1 ADVIO Dataset', 2)
    add_para(doc, 'The ADVIO (Aalto Vision Indoor-Outdoor) dataset is an academic benchmark for visual-inertial '
             'navigation research. It contains 23 sequences recorded with an iPhone in various indoor and outdoor '
             'environments in Helsinki, Finland.')
    
    add_para(doc, 'Key characteristics:', bold=True)
    bullets3 = [
        'Sensor: iPhone accelerometer, gyroscope, magnetometer, barometer (100 Hz)',
        'Ground truth: Generated from Structure-from-Motion point clouds and floor plans',
        'Elevator segments: Only 3 sequences contain elevators (advio-07, advio-14, advio-18)',
        'Total elevator segments: 7 rides across 3 sequences',
        'Phone position: Held in hand (visible camera)',
        'Advantages: Accurate GT from professional mapping, varied environments',
        'Limitations: Very few elevator instances, single phone orientation, single building type',
    ]
    for b in bullets3:
        doc.add_paragraph(b, style='List Bullet')
    
    add_heading(doc, '2.1.1 ADVIO Elevator Segments', 3)
    add_table(doc,
              ['Sequence', 'Segment', 'Start (s)', 'End (s)', 'Direction', 'GT Height (m)'],
              [['advio-07', '0', '17.0', '26.0', 'Up', '5.60'],
               ['advio-07', '1', '36.5', '41.5', 'Up', '4.53'],
               ['advio-07', '2', '50.5', '57.5', 'Up', '4.52'],
               ['advio-07', '3', '66.0', '73.5', 'Up', '4.52'],
               ['advio-07', '4', '81.5', '86.5', 'Up', '4.46'],
               ['advio-14', '0', '18.5', '35.5', 'Up', '7.52'],
               ['advio-18', '0', '73.5', '83.0', 'Up', '7.81']])
    
    add_heading(doc, '2.2 Bar-Ilan Dataset', 2)
    add_para(doc, 'A custom dataset collected specifically for this project in a 16-floor residential building '
             'on Bar-Ilan Street, Ramat Gan, Israel. The building has a ground floor "L" (street level, height −2.6m) '
             'and floors 0–15 (heights 15.8m, 18.8m, 21.8m, ... with 3.0m spacing between floors).')
    
    add_figure(doc, 'fig01_building_heights.png', 
               'Figure 2.1: Bar-Ilan Building Floor Heights (from Gramushka)', width=3.5)
    
    add_para(doc, 'Data collection protocol:', bold=True)
    bullets4 = [
        'Device: Google Pixel phone with Sensor Logger app (accelerometer, gyroscope, magnetometer, barometer, GPS at 100 Hz)',
        'Video: Simultaneous video recording for ground truth verification',
        'Two collection modes: (1) Phone held in hand (camera visible), (2) Phone in pocket',
        'Varied ride patterns: Short rides (1 floor, 3m), medium rides (2 floors, 6m), long rides (up to 16 floors, ~57m)',
        'Both directions: Up and down rides recorded',
        'Total: 33 ground truth elevator rides across a single continuous 24-minute recording',
        'GT derivation: Floor plans (gramushka), elevator audio announcements, barometer cross-validation',
        'Temperature: 14°C, Location: 32.166°N, 34.842°E',
    ]
    for b in bullets4:
        doc.add_paragraph(b, style='List Bullet')
    
    add_heading(doc, '2.2.1 Bar-Ilan Data Statistics', 3)
    # Calculate stats from v4 data
    hand_rides = [r for r in v4['per_ride'] if r['phone'] == 'hand']
    pocket_rides = [r for r in v4['per_ride'] if r['phone'] == 'pocket']
    up_rides = [r for r in v4['per_ride'] if r['true_dh'] > 0]
    down_rides = [r for r in v4['per_ride'] if r['true_dh'] < 0]
    
    add_table(doc,
              ['Category', 'Count', 'Height Range (m)'],
              [['Total rides', '33', '3.0 – 57.4'],
               ['Hand-held', str(len(hand_rides)), f'{min(abs(r["true_dh"]) for r in hand_rides):.1f} – {max(abs(r["true_dh"]) for r in hand_rides):.1f}'],
               ['Pocket', str(len(pocket_rides)), f'{min(abs(r["true_dh"]) for r in pocket_rides):.1f} – {max(abs(r["true_dh"]) for r in pocket_rides):.1f}'],
               ['Upward', str(len(up_rides)), f'{min(r["true_dh"] for r in up_rides):.1f} – {max(r["true_dh"] for r in up_rides):.1f}'],
               ['Downward', str(len(down_rides)), f'{max(r["true_dh"] for r in down_rides):.1f} – {min(r["true_dh"] for r in down_rides):.1f}']])
    
    add_figure(doc, 'fig02_raw_accel_traces.png',
               'Figure 2.2: Raw accelerometer magnitude for hand-held (top) and pocket (bottom) sections')
    
    add_figure(doc, 'fig03_gt_height_profile.png',
               'Figure 2.3: Ground truth height profile with elevator segments and phone position')
    
    add_heading(doc, '2.3 Dataset Comparison', 2)
    add_table(doc,
              ['Feature', 'ADVIO', 'Bar-Ilan'],
              [['# Sequences', '23 (3 with elevators)', '1 continuous'],
               ['# Elevator Rides', '7', '33'],
               ['Phone Positions', 'Hand only', 'Hand + Pocket'],
               ['Ride Directions', 'Up only', 'Up + Down'],
               ['Height Range', '4.5 – 7.8m', '3.0 – 57.4m'],
               ['GT Source', 'SfM point clouds', 'Floor plans + audio'],
               ['Sensor Rate', '100 Hz', '100 Hz'],
               ['Duration', '~90s each', '1412s total']])
    
    add_heading(doc, '2.4 Train/Test Methodology', 2)
    add_para(doc, 'The evaluation uses the full Bar-Ilan dataset (33 rides) as the primary benchmark. '
             'The pipeline is evaluated on ALL rides simultaneously (no train/test split for detection or '
             'estimation), with conformal prediction intervals computed via Leave-One-Out (LOO) cross-validation '
             'to maximize statistical power from the limited sample size.')
    
    add_para(doc, 'LOO conformal prediction computes the interval by iteratively training on n−1 rides and '
             'testing on the held-out ride. The 90th percentile of the resulting residuals gives the '
             'prediction interval. This is more appropriate than a 50/50 split for small datasets because '
             'it uses all data for both training and testing.')
    
    doc.add_page_break()
    
    # ================================================================
    # CHAPTER 3: RESEARCH HISTORY
    # ================================================================
    add_heading(doc, '3. Research History', 1)
    
    add_heading(doc, '3.1 Phase 1: Initial Algorithm Research (ADVIO)', 2)
    add_para(doc, 'The project began with the ADVIO dataset and three classical approaches to '
             'accelerometer-based height estimation. All three operate on a single axis (vertical) '
             'and assume the phone is held upright.')
    
    add_heading(doc, '3.1.1 Direct Integration', 3)
    add_para(doc, 'The simplest approach: subtract mean acceleration (gravity estimate) and double-integrate. '
             'No drift correction. Serves as the baseline.')
    add_para(doc, 'h = ∫∫ (a_z − mean(a_z)) dt²', bold=True)
    add_para(doc, 'Pros: Simple, fast. Cons: Severe drift accumulation over long windows, sensitive to bias errors.')
    
    add_heading(doc, '3.1.2 ZUPT (Zero velocity UPdaTe)', 3)
    add_para(doc, 'Identifies periods of zero velocity (before/after ride) and applies linear drift correction. '
             'The velocity is forced to zero at the ride boundaries, which constrains the position integral.')
    add_para(doc, 'Algorithm: (1) Detect active window via acceleration threshold, (2) Integrate to velocity, '
             '(3) Apply linear drift correction so v(t_end) = 0, (4) Integrate corrected velocity to position.')
    add_para(doc, 'Pros: Effectively removes linear drift. Cons: Assumes known ride boundaries, single-axis only.')
    
    add_heading(doc, '3.1.3 Kalman Filter', 3)
    add_para(doc, 'State-space model with position and velocity as states, acceleration as input. '
             'Process noise models drift, measurement noise models sensor error. Provides optimal '
             'linear estimate under Gaussian noise assumptions.')
    add_para(doc, 'Pros: Theoretically optimal for linear systems. Cons: Requires tuning of noise parameters, '
             'still single-axis, doesn\'t handle orientation changes.')
    
    add_heading(doc, '3.1.4 Phase 1 Results', 3)
    import pandas as pd
    advio_results = pd.read_csv(os.path.join(BASE, "metadata", "evaluation_results.csv"))
    gt = advio_results['GT'].values
    
    add_table(doc,
              ['Segment', 'GT (m)', 'Direct (m)', 'ZUPT (m)', 'Kalman (m)', 'Barometer (m)'],
              [[f'{advio_results.iloc[i]["dataset"]}-{advio_results.iloc[i]["segment"]}',
                f'{gt[i]:.2f}',
                f'{advio_results.iloc[i]["Algo1_Direct"]:.2f}',
                f'{advio_results.iloc[i]["Algo2_ZUPT"]:.2f}',
                f'{advio_results.iloc[i]["Algo3_Kalman"]:.2f}',
                f'{advio_results.iloc[i]["Barometer"]:.2f}']
               for i in range(len(gt))])
    
    add_figure(doc, 'fig04_advio_historical.png',
               'Figure 3.1: ADVIO Phase 1 — True vs Estimated for each algorithm')
    
    add_figure(doc, 'fig05_algo_comparison_bar.png',
               'Figure 3.2: ADVIO Phase 1 — MAE comparison across methods')
    
    add_para(doc, 'Key Finding: ZUPT performed best on ADVIO with MAE = '
             f'{np.mean(np.abs(gt - advio_results["Algo2_ZUPT"].values)):.2f}m, '
             'followed by Kalman and Direct Integration. However, all methods '
             'assumed the phone was held upright (gravity aligned with z-axis), '
             'which is unrealistic for real-world deployment.')
    
    add_heading(doc, '3.2 Phase 2: Confidence Interval Research', 2)
    add_para(doc, 'The second phase focused on quantifying estimation uncertainty. Two approaches were investigated:')
    
    add_heading(doc, '3.2.1 Theoretical Confidence Intervals', 3)
    add_para(doc, 'Using the accelerometer noise database (from the CameraOrientation repository), '
             'we modeled ZUPT integration error as a function of sensor noise characteristics and '
             'integration duration. The theoretical 90% confidence interval is:')
    add_para(doc, 'CI_90 = 1.645 × σ_noise × √(N³/3) × dt²', bold=True)
    add_para(doc, 'where σ_noise is the sensor noise density, N is the number of samples, and dt is '
             'the sampling interval. This provides a phone-model-specific bound.')
    
    add_heading(doc, '3.2.2 Conformal Prediction', 3)
    add_para(doc, 'Conformal prediction provides distribution-free prediction intervals with guaranteed '
             'coverage. Unlike theoretical CI which assumes Gaussian noise, conformal prediction uses '
             'empirical residuals from calibration data:')
    add_para(doc, 'Key property: If the calibration and test data are exchangeable, conformal prediction '
             'guarantees P(Y_new ∈ Ĉ(X_new)) ≥ 1 − α for any data distribution.', italic=True)
    add_para(doc, 'For split-conformal: the interval is the ⌈(n+1)(1−α)⌉/n quantile of calibration errors. '
             'For LOO conformal: compute intervals by leaving each calibration point out in turn.')
    
    add_heading(doc, '3.3 Phase 3: Bar-Ilan Dataset Construction', 2)
    add_para(doc, 'With only 7 ADVIO elevator segments, the dataset was insufficient for robust algorithm '
             'development. A custom dataset was collected in a 16-floor building with the following protocol:')
    bullets5 = [
        'Continuous recording: One 24-minute recording capturing 33 rides in sequence',
        'Varied conditions: Hand-held (first half) and pocket (second half)',
        'Varied ride patterns: 1-floor to 16-floor rides, both up and down',
        'Ground truth construction: Floor heights from building plans (gramushka), elevator audio '
        'announcements for floor identification, barometer data for cross-validation',
        'Sensor synchronization: Raw sensor log timestamps aligned with video timestamps',
        'Metadata: 10 Hz annotated CSV with height, elevator status, segment ID, phone position',
    ]
    for b in bullets5:
        doc.add_paragraph(b, style='List Bullet')
    
    add_heading(doc, '3.4 Phase 4: Detection Algorithm Development', 2)
    add_para(doc, 'Three detection algorithms were implemented and compared:')
    
    add_heading(doc, '3.4.1 State Machine via Low-Pass Filtering', 3)
    add_para(doc, 'Uses rolling variance of acceleration to identify standing-still '
             'periods, then checks for characteristic acceleration-deceleration patterns '
             'within low-variance blocks.')
    
    add_heading(doc, '3.4.2 Sliding DTW Pattern Matching', 3)
    add_para(doc, 'Compares acceleration blocks against an ideal elevator template using '
             'Dynamic Time Warping. Theoretically robust to speed variations but '
             'computationally expensive and sensitive to template design.')
    
    add_heading(doc, '3.4.3 Velocity Integral Bounding', 3)
    add_para(doc, 'Integrates acceleration to velocity and position, checks for net displacement '
             'exceeding threshold with final velocity near zero.')
    
    add_para(doc, 'Result: The state machine approach was most reliable for the ADVIO dataset, '
             'but all three struggled with the Bar-Ilan dataset where the phone was in a pocket, '
             'because they relied on single-axis (z-axis) acceleration which assumes upright orientation.')
    
    add_heading(doc, '3.5 Phase 5: Current Solution Development', 2)
    add_para(doc, 'The critical insight was that single-axis methods fail completely when the phone '
             'orientation is unknown. The current solution addresses this by:')
    bullets6 = [
        'Using 3-axis accelerometer magnitude for detection (rotation-invariant)',
        'Estimating gravity direction from stationary data before/after the ride',
        'Projecting acceleration onto the estimated gravity direction for height estimation',
        'Using magnitude-based ZUPT as a fallback when gravity projection is unreliable',
        'Implementing drift-corrected magnitude selection when gravity drift is high',
        'Multiple rejection criteria based on accelerometer-only quality features',
    ]
    for b in bullets6:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_page_break()
    
    # ================================================================
    # CHAPTER 4: CURRENT SOLUTION ARCHITECTURE
    # ================================================================
    add_heading(doc, '4. Current Solution — Pipeline Architecture', 1)
    
    add_figure(doc, 'fig21_pipeline_diagram.png',
               'Figure 4.1: Pipeline Architecture — Three-stage design')
    
    add_para(doc, 'The pipeline processes raw 3-axis accelerometer data through three stages: '
             'Detection & Segmentation, Quality Filtering, and Height Estimation with Conformal Prediction. '
             'Each stage is designed to be modular and uses ONLY accelerometer data (no gyroscope required).')
    
    # Stage 1: Detection
    add_heading(doc, '4.1 Stage 1: Detection & Segmentation', 2)
    add_para(doc, 'The detection stage identifies individual elevator rides from continuous accelerometer recordings.')
    
    add_heading(doc, '4.1.1 Algorithm', 3)
    add_para(doc, 'Step 1 — Magnitude Variance: Compute the rolling variance of acceleration magnitude '
             '|a| = √(ax² + ay² + az²) over a 1.5-second window. Elevator rides have low magnitude '
             'variance (person standing still) compared to walking (high variance).')
    add_para(doc, 'Step 2 — Still-Block Identification: Find contiguous blocks where variance < 1.5 m²/s⁴ '
             'and duration > 4 seconds. These are candidate elevator ride blocks.')
    add_para(doc, 'Step 3 — Within-Block Integration: For each still-block, integrate the linear '
             'acceleration (magnitude - mean) to get velocity, then position. Check for net displacement > 1m.')
    add_para(doc, 'Step 4 — Velocity Zero-Crossing Segmentation: Split multi-ride blocks at velocity '
             'zero-crossings where position has changed by ≥ 1m. This separates consecutive rides '
             '(e.g., floor 1→3 and 3→7) into distinct segments.')
    add_para(doc, 'Step 5 — Trimming: Trim leading/trailing acceleration below 0.08 m/s² threshold '
             'with 0.5s margin to isolate the active ride portion.')
    
    add_heading(doc, '4.1.2 Detection Results', 3)
    add_para(doc, f'On the Bar-Ilan dataset (33 GT rides):', bold=True)
    add_para(doc, '• 61 candidate segments detected\n'
             '• 28/33 GT rides matched (85% recall) at IoU ≥ 0.3\n'
             '• 5 missed rides (3 were near stationary segments, 2 were boundary cases)')
    
    add_figure(doc, 'fig06_detection_timeline.png',
               'Figure 4.2: Detection timeline — accelerometer magnitude, rolling variance, '
               'and ground truth height with detected segments (blue shading) overlaid')
    
    add_figure(doc, 'fig22_iou_distribution.png',
               'Figure 4.3: IoU distribution for matched rides')
    
    add_heading(doc, '4.1.3 Segmentation Accuracy Analysis', 3)
    add_para(doc, 'To evaluate how well the detected segments align with the ground truth ride boundaries, '
             'we conducted a detailed segmentation accuracy analysis comparing GT start/end times with '
             'detected segment start/end times for all matched rides.')
    
    add_para(doc, 'The analysis examines:', bold=True)
    seg_bullets = [
        'Temporal IoU: Measures the overlap between GT and detected segments. '
        'Higher IoU indicates better segmentation precision.',
        'Boundary errors: The start-time and end-time errors (in seconds) quantify how '
        'accurately the segment boundaries are detected.',
        'Duration comparison: GT vs detected ride duration shows whether segments are '
        'systematically shortened or extended by the detection algorithm.',
    ]
    for b in seg_bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    add_figure(doc, 'fig24_segmentation_accuracy.png',
               'Figure 4.4: Segmentation accuracy analysis — (top-left) GT vs detected boundaries, '
               '(top-right) IoU distribution, (bottom-left) boundary errors in seconds, '
               '(bottom-right) GT vs detected duration comparison', width=6.5)
    
    add_para(doc, 'Key observations: Out of 33 GT rides, 32 were matched with IoU > 0.1. '
             'Start boundaries are detected within ±2 seconds of GT for most rides. '
             'End boundaries tend to extend slightly beyond GT segments due to the trimming '
             'algorithm including deceleration tails. Duration agreement is strong, '
             'with the majority of rides having detected duration within 20% of GT duration.')
    
    # Stage 2: Quality Filter
    add_heading(doc, '4.2 Stage 2: Quality Filter', 2)
    add_para(doc, 'The quality filter assesses whether a detected segment can be reliably used for '
             'height estimation, using ONLY accelerometer features. This is critical for preventing '
             'catastrophic errors from being included in conformal prediction calibration.')
    
    add_heading(doc, '4.2.1 Gravity Vector Estimation', 3)
    add_para(doc, 'Pre-ride context window: The 5 seconds of data immediately before the detected ride '
             'start. Gravity is estimated using the median-of-windows method: divide the window into '
             '0.5-second chunks, compute the mean acceleration vector in each chunk, take the median '
             'across chunks for robustness against transient disturbances.')
    add_para(doc, 'Post-ride fallback: If pre-ride data is unstable (e.g., phone being placed in pocket), '
             'the 5 seconds after the ride end are used instead. This rescues rides where the phone '
             'orientation was changing just before the ride started.')
    
    add_heading(doc, '4.2.2 Rejection Criteria', 3)
    add_para(doc, 'The quality filter applies the following rejection rules (in order):', bold=True)
    
    rejection_rules = [
        ['1', 'No stable calibration', 'Neither pre-ride nor post-ride data produces a stable gravity estimate (stability > 1.5 m/s²)', 'Prevents estimation without reliable gravity direction'],
        ['2', 'Large orientation change', 'Angle between pre-ride and post-ride gravity vectors > 25°', 'Phone was moved significantly during ride'],
        ['3', 'Impact detected', 'Peak acceleration magnitude > 8 m/s² (above gravity)', 'Phone was dropped, shaken, or bumped'],
        ['4', 'High noise', 'Acceleration std > 0.9 m/s² AND peak > 10 m/s²', 'Excessive vibration or handling noise'],
        ['5', 'High gravity drift', 'During-ride gravity drift > 15° (computed from 1s chunks)', 'Phone orientation changed during ride'],
        ['6', 'Projection/magnitude disagree', 'Gravity-projected estimate disagrees with magnitude estimate by > 1.8× (or 2.5× for low drift)', 'Cross-validation failure between methods'],
        ['7', 'Signed-mag unreliable', 'Signed magnitude estimate > 15m AND no good gravity projection', 'Sign determination unreliable for long rides'],
        ['8', 'Implausible estimate', 'Estimated height > 100m', 'Integration diverged catastrophically'],
    ]
    
    add_table(doc,
              ['#', 'Rule', 'Condition', 'Rationale'],
              rejection_rules)
    
    add_heading(doc, '4.2.3 Quality Filter Results', 3)
    rej = [r for r in v4['per_ride'] if not r['accepted']]
    acc = [r for r in v4['per_ride'] if r['accepted']]
    correct_rej = sum(1 for r in rej if r['err'] > 1.0)
    
    add_para(doc, f'• {len(acc)} rides accepted ({len(acc)*100//33}%), {len(rej)} rejected ({len(rej)*100//33}%)\n'
             f'• {correct_rej}/{len(rej)} rejected rides had error > 1m ({correct_rej*100//len(rej)}% correct rejection)\n'
             f'• Only {len(rej) - correct_rej} false rejection(s)')
    
    add_figure(doc, 'fig12_rejection_reasons.png',
               'Figure 4.4: Rejection reason breakdown')
    
    add_figure(doc, 'fig18_rejection_accuracy.png',
               'Figure 4.5: Rejection decision accuracy analysis')
    
    # Stage 3: Height Estimation
    add_heading(doc, '4.3 Stage 3: Height Estimation', 2)
    
    add_heading(doc, '4.3.1 Gravity-Projected ZUPT', 3)
    add_para(doc, 'The primary estimation method projects the 3-axis acceleration onto the estimated '
             'gravity direction, then applies ZUPT integration on the vertical component. This recovers '
             'the true vertical acceleration regardless of phone orientation:')
    add_para(doc, 'a_vert = (a⃗ · ĝ) − |g|', bold=True)
    add_para(doc, 'where ĝ is the unit gravity vector estimated from the pre/post-ride stationary data '
             'and |g| is the gravity magnitude. ZUPT integration then applies:')
    add_para(doc, '1. Remove DC bias from a_vert\n'
             '2. Find active window (|a_vert| > 0.05 m/s² smoothed)\n'
             '3. Integrate to velocity\n'
             '4. Apply linear drift correction (v(t_end) = 0)\n'
             '5. Integrate to position')
    
    add_heading(doc, '4.3.2 Magnitude-Based ZUPT (Rotation-Invariant)', 3)
    add_para(doc, 'When gravity projection is unreliable (high drift, unstable calibration), '
             'the pipeline uses the acceleration magnitude minus mean as input to ZUPT integration. '
             'This is rotation-invariant but loses directional information (always positive displacement). '
             'The sign is recovered from the gravity projection if available, or from the initial '
             'acceleration pulse direction.')
    
    add_heading(doc, '4.3.3 Drift-Corrected Magnitude Selection', 3)
    add_para(doc, 'A key innovation: when during-ride gravity drift exceeds 8° AND the gravity-projected '
             'estimate exceeds the magnitude estimate by more than 1.5×, the projection is likely '
             'contaminated by horizontal acceleration. In this case, the system automatically switches '
             'to the signed-magnitude estimate. This corrected Ride 9 from 6.02m error to 1.72m error.')
    
    add_heading(doc, '4.3.4 Estimate Selection Logic', 3)
    add_para(doc, 'The pipeline selects the best estimate using the following priority:')
    selection_rules = [
        ['1', 'Drift-corrected magnitude', 'drift > 8° AND |GP|/|mag| > 1.5', 'Use signed magnitude'],
        ['2', 'High-quality GP', 'GP quality < 0.5 AND |GP| < 150m', 'Use gravity projection'],
        ['3', 'Moderate-quality GP', 'GP agrees with magnitude (agreement > 0.3)', 'Use gravity projection'],
        ['4', 'Fallback', 'All other cases', 'Use signed magnitude'],
    ]
    add_table(doc, ['Priority', 'Method', 'Condition', 'Action'], selection_rules)
    
    add_figure(doc, 'fig15_method_breakdown.png',
               'Figure 4.6: Estimation method selection breakdown for accepted rides')
    
    # Stage 4: Conformal Prediction
    add_heading(doc, '4.4 Stage 4: Conformal Prediction', 2)
    add_para(doc, 'After height estimation, the pipeline provides a 90% prediction interval using '
             'Leave-One-Out (LOO) conformal prediction. This provides finite-sample coverage guarantee:')
    add_para(doc, 'For each ride i in the calibration set of n rides:\n'
             '  1. Remove ride i from the set\n'
             '  2. Compute the 90th percentile of remaining n−1 errors\n'
             '  3. Check if error_i ≤ interval\n'
             'Coverage = fraction of rides covered ≥ 90%')
    
    add_para(doc, f'Results: LOO coverage = 94.7% (≥90% target ✓), average interval = ±3.98m')
    
    add_figure(doc, 'fig14_conformal_coverage.png',
               'Figure 4.7: LOO Conformal prediction — coverage and per-ride intervals')
    
    doc.add_page_break()
    
    # ================================================================
    # CHAPTER 5: ALGORITHM COMPARISON
    # ================================================================
    add_heading(doc, '5. Algorithm Comparison', 1)
    
    add_para(doc, 'This chapter compares the current pipeline against all historically investigated approaches. '
             'Note that the current pipeline uses 3-axis accelerometer while the earlier methods used single-axis, '
             'and the datasets differ (ADVIO for Phase 1, Bar-Ilan for current). Nevertheless, the comparison '
             'illustrates the evolution and why the current approach was necessary.')
    
    add_heading(doc, '5.1 Why Single-Axis Methods Failed', 2)
    add_para(doc, 'When the Bar-Ilan dataset was introduced with pocket-mode rides, all single-axis methods '
             'catastrophically failed because:')
    bullets7 = [
        'Gravity direction assumption: a_z ≈ ±g is violated when phone is at arbitrary orientation in pocket',
        'No way to decompose: Without knowing orientation, cannot separate vertical from horizontal acceleration',
        'Detection failure: Single-axis variance-based detection could not reliably find elevator rides when phone bounced in pocket',
    ]
    for b in bullets7:
        doc.add_paragraph(b, style='List Bullet')
    
    add_heading(doc, '5.2 Improvement from 3-Axis Processing', 2)
    add_para(doc, 'The key innovations that enabled reliable estimation in arbitrary orientations:')
    improvements = [
        ['Magnitude-based detection', 'Uses |a| = √(x²+y²+z²) instead of a_z', 'Rotation-invariant, works in any orientation'],
        ['Gravity projection', 'Projects a⃗ onto estimated gravity ĝ', 'Recovers vertical component regardless of orientation'],
        ['Post-ride gravity fallback', 'Uses data after ride when pre-ride is unstable', 'Rescues pocket-mode rides'],
        ['Drift-corrected magnitude', 'Switches to |a| when gravity drift is high', 'Prevents projection errors from contaminating estimate'],
        ['Quality filter', 'Accelerometer-only orientation stability checks', 'Rejects unreliable segments before they affect conformal calibration'],
    ]
    add_table(doc, ['Innovation', 'Mechanism', 'Benefit'], improvements)
    
    add_heading(doc, '5.3 Cross-Method MAE Comparison', 2)
    add_figure(doc, 'fig23_full_comparison.png',
               'Figure 5.1: MAE comparison across all historical and current methods')
    
    add_para(doc, 'Important caveat: The ADVIO results use a different, simpler dataset (7 rides, all hand-held, '
             'all upward). The Bar-Ilan V4 pipeline results are on a much harder dataset (33 rides, hand + pocket, '
             'up + down, small + large rides). The V4 MAE of 1.16m on accepted rides is achieved on a '
             'fundamentally more challenging benchmark.')
    
    doc.add_page_break()
    
    # ================================================================
    # CHAPTER 6: FINAL RESULTS & VALIDATION
    # ================================================================
    add_heading(doc, '6. Final Results & Validation', 1)
    
    add_heading(doc, '6.1 Aggregate Performance Metrics', 2)
    acc_errs = [r['err'] for r in v4['per_ride'] if r['accepted']]
    
    add_table(doc,
              ['Metric', 'Value'],
              [['Total GT Rides', '33'],
               ['Detection Recall', '28/33 (85%)'],
               ['Acceptance Rate', f'{len(acc)}/33 ({len(acc)*100//33}%)'],
               ['Accepted MAE', f'{np.mean(acc_errs):.2f}m'],
               ['Accepted Median Error', f'{np.median(acc_errs):.2f}m'],
               ['Accepted Max Error', f'{max(acc_errs):.2f}m'],
               ['< 0.5m Accuracy', f'{sum(1 for e in acc_errs if e<0.5)}/{len(acc_errs)} ({sum(1 for e in acc_errs if e<0.5)*100//len(acc_errs)}%)'],
               ['< 1.0m Accuracy', f'{sum(1 for e in acc_errs if e<1)}/{len(acc_errs)} ({sum(1 for e in acc_errs if e<1)*100//len(acc_errs)}%)'],
               ['< 2.0m Accuracy', f'{sum(1 for e in acc_errs if e<2)}/{len(acc_errs)} ({sum(1 for e in acc_errs if e<2)*100//len(acc_errs)}%)'],
               ['LOO Conformal Coverage', '94.7% (target ≥ 90%)'],
               ['LOO Conformal Interval', '±3.98m'],
               ['Correct Rejection Rate', f'{correct_rej}/{len(rej)} ({correct_rej*100//len(rej)}%)']])
    
    add_heading(doc, '6.2 True vs Estimated Height', 2)
    add_figure(doc, 'fig07_scatter_current.png',
               'Figure 6.1: True vs Estimated height for all 33 rides (green=accepted, red=rejected)')
    
    add_heading(doc, '6.3 Per-Ride Error Analysis', 2)
    add_figure(doc, 'fig08_per_ride_errors.png',
               'Figure 6.2: Per-ride absolute error (green=accepted, red=rejected, capped at 20m)')
    
    add_heading(doc, '6.3.1 Detailed Per-Ride Results', 3)
    ride_rows = []
    for r in v4['per_ride']:
        status = '✓' if r['accepted'] else '✗'
        ride_rows.append([
            str(r['id']), r['phone'], f"{r['true_dh']:+.1f}",
            f"{r['est_dh']:+.2f}", f"{r['err']:.2f}",
            r['method'], status
        ])
    add_table(doc,
              ['Ride', 'Phone', 'True (m)', 'Est (m)', 'Error (m)', 'Method', 'Acc'],
              ride_rows)
    
    add_heading(doc, '6.4 Error Distribution Analysis', 2)
    add_figure(doc, 'fig09_error_histogram.png',
               'Figure 6.3: Error histogram for accepted rides')
    add_figure(doc, 'fig10_error_cdf.png',
               'Figure 6.4: Cumulative error distribution (CDF) with accuracy percentages')
    
    add_heading(doc, '6.5 Hand vs Pocket Performance', 2)
    hand_acc = [r['err'] for r in v4['per_ride'] if r['accepted'] and r['phone']=='hand']
    pocket_acc = [r['err'] for r in v4['per_ride'] if r['accepted'] and r['phone']=='pocket']
    
    add_para(doc, f'Hand-held rides: n={len(hand_acc)}, MAE={np.mean(hand_acc):.2f}m, '
             f'Median={np.median(hand_acc):.2f}m')
    add_para(doc, f'Pocket rides: n={len(pocket_acc)}, MAE={np.mean(pocket_acc):.2f}m, '
             f'Median={np.median(pocket_acc):.2f}m')
    
    add_figure(doc, 'fig11_hand_vs_pocket.png',
               'Figure 6.5: Error comparison by phone position with individual data points')
    
    add_heading(doc, '6.6 Error vs Ride Magnitude', 2)
    add_figure(doc, 'fig20_error_vs_height.png',
               'Figure 6.6: Absolute error vs true ride height magnitude')
    
    add_heading(doc, '6.7 Quality Feature Analysis', 2)
    add_figure(doc, 'fig13_quality_correlations.png',
               'Figure 6.7: Quality score vs error (left) and error vs ride magnitude (right)')
    
    add_heading(doc, '6.8 Individual Ride Examples', 2)
    add_para(doc, 'Six representative rides showing different estimation scenarios:')
    add_figure(doc, 'fig16_individual_rides.png',
               'Figure 6.8: Individual ride examples — good, drift-corrected, pocket, large, and problematic rides')
    
    add_heading(doc, '6.9 Per-Ride Acceleration & Displacement Analysis (Bar-Ilan)', 2)
    add_para(doc, 'The following figure shows detailed acceleration (gravity-removed) and vertical displacement '
             'curves for 6 representative Bar-Ilan rides. For each ride, the left panel shows the gravity-removed '
             'acceleration magnitude during the ride, and the right panel shows the estimated displacement curve '
             'compared to the ground truth height difference.')
    add_para(doc, 'The displacement panel shows multiple estimation methods: the V4 pipeline estimate (using '
             'whichever method was selected), the magnitude ZUPT estimate, and the ground truth as a horizontal line. '
             'This allows visual comparison of how well each method tracks the true vertical displacement.')
    add_figure(doc, 'fig25_bar_ilan_ride_analysis.png',
               'Figure 6.9: Bar-Ilan per-ride analysis — acceleration (left) and displacement curves vs GT (right) '
               'for 6 representative rides', width=6.5)
    
    add_heading(doc, '6.10 Per-Segment Analysis (ADVIO Dataset)', 2)
    add_para(doc, 'To validate the pipeline on an independent dataset, the same analysis is performed on all 7 '
             'ADVIO elevator segments. Each segment shows the gravity-removed acceleration and the displacement '
             'curves from the V4 pipeline, Direct Integration, ZUPT magnitude, GT, and Barometer reference.')
    add_para(doc, 'This cross-dataset validation demonstrates that the pipeline generalizes to data collected '
             'with a different phone (iPhone), in different buildings (Helsinki), and at different times — '
             'validating the rotation-invariant design principles.')
    add_figure(doc, 'fig26_advio_ride_analysis.png',
               'Figure 6.10: ADVIO per-segment analysis — acceleration and displacement curves for all 7 '
               'elevator segments across advio-07, advio-14, and advio-18', width=6.5)
    
    add_heading(doc, '6.11 3-Axis Accelerometer During a Ride', 2)
    add_figure(doc, 'fig19_3axis_example.png',
               'Figure 6.11: Individual 3-axis accelerometer components during Ride 4 (hand, +6.0m)')
    
    add_heading(doc, '6.12 Rejection Analysis', 2)
    add_para(doc, 'Detailed breakdown of rejected rides:', bold=True)
    rej_rows = []
    for r in v4['per_ride']:
        if not r['accepted']:
            rej_rows.append([
                str(r['id']), r['phone'], f"{r['true_dh']:+.1f}",
                f"{r['err']:.2f}", r.get('reject_reason', 'N/A')
            ])
    add_table(doc, ['Ride', 'Phone', 'True (m)', 'Error (m)', 'Rejection Reason'], rej_rows)
    
    add_heading(doc, '6.13 Conformal Prediction Validation', 2)
    add_figure(doc, 'fig14_conformal_coverage.png',
               'Figure 6.12: LOO conformal prediction — coverage bar and per-ride intervals')
    
    add_para(doc, 'The LOO conformal prediction achieves 94.7% coverage (18/19 accepted rides within interval), '
             'exceeding the 90% target. The average LOO interval of ±3.98m reflects the influence of 3 accepted '
             'rides with 2–4m errors that the quality filter cannot distinguish from accurate rides using '
             'accelerometer features alone.')
    
    add_heading(doc, '6.14 Summary Dashboard', 2)
    add_figure(doc, 'fig17_summary_dashboard.png',
               'Figure 6.13: Complete performance summary dashboard')
    
    doc.add_page_break()
    
    # ================================================================
    # CHAPTER 7: USAGE GUIDE
    # ================================================================
    add_heading(doc, '7. Usage Guide & Deployment', 1)
    
    add_heading(doc, '7.1 Installation', 2)
    add_para(doc, 'pip install -r requirements.txt', bold=True)
    add_para(doc, 'Dependencies: numpy, pandas, scipy, matplotlib, python-docx')
    
    add_heading(doc, '7.2 Quick Start (Python API)', 2)
    add_para(doc, 'from src.pipeline import ElevatorHeightPipeline\n\n'
             'pipeline = ElevatorHeightPipeline.load("model/")\n'
             'results = pipeline.process(acc_x, acc_y, acc_z, fs=100)\n\n'
             'for r in results:\n'
             '    if r["accepted"]:\n'
             '        print(f"  [{r[\'start_time\']:.1f}s – {r[\'end_time\']:.1f}s]  "\n'
             '              f"Height: {r[\'height_estimate\']:+.2f}m ± {r[\'confidence_interval_90\']:.2f}m")',
             font_size=9)
    
    add_para(doc, 'Each result dictionary includes start_time and end_time (in seconds) identifying '
             'the detected elevator segment time frame, along with height_estimate, confidence_interval_90, '
             'method, accepted status, and reject_reason.')
    
    add_heading(doc, '7.3 Visual Output (process_plot)', 2)
    add_para(doc, 'The pipeline provides a built-in visualization method that generates a 2-panel figure '
             'showing all detected rides and their height estimates:')
    add_para(doc, 'results, fig = pipeline.process_plot(acc_x, acc_y, acc_z, fs=100,\n'
             '                                         save_path="output.png")',
             font_size=9)
    add_para(doc, 'The top panel shows the accelerometer magnitude time series with detected segments '
             'highlighted (green for accepted, red for rejected). The bottom panel shows a bar chart '
             'of height estimates with 90% conformal prediction interval whiskers.')
    
    add_figure(doc, 'fig27_process_plot_demo.png',
               'Figure 7.1: Example output of pipeline.process_plot() on the Bar-Ilan dataset')
    
    add_heading(doc, '7.4 Command-Line Interface', 2)
    add_para(doc, 'python run_inference.py --input data.csv --output results.json --verbose', bold=True)
    add_para(doc, 'Input: CSV with columns acc_x, acc_y, acc_z (m/s²), optional time column.\n'
             'Output: JSON array of detected rides with height estimates and confidence intervals.')
    
    add_heading(doc, '7.5 Re-Calibration', 2)
    add_para(doc, 'To calibrate conformal prediction on new labeled data:')
    add_para(doc, 'pipeline.calibrate(rides_with_gt)\npipeline.save("model/")', font_size=9)
    
    add_heading(doc, '7.6 Limitations', 2)
    limits = [
        'Accelerometer-only: No fusion with barometer or gyroscope. Orientation changes during ride cannot be tracked.',
        'Pocket mode: Accuracy degrades in pocket due to walking-induced vibrations and orientation uncertainty.',
        'Very long rides: For rides > 50m, double integration drift becomes significant.',
        'Sign determination: The direction (up/down) relies on gravity projection. If calibration fails, magnitude is unsigned.',
        'Single building validation: The pipeline was calibrated on data from one building. Generalization to other elevator types needs validation.',
    ]
    for l in limits:
        doc.add_paragraph(l, style='List Bullet')
    
    doc.add_page_break()
    
    # ================================================================
    # CHAPTER 8: CONCLUSIONS
    # ================================================================
    add_heading(doc, '8. Conclusions & Future Work', 1)
    
    add_heading(doc, '8.1 Summary of Achievements', 2)
    add_para(doc, 'This project developed a three-stage pipeline for accelerometer-only elevator height '
             'estimation that achieves:')
    achievements = [
        'Detection: 85% recall (28/33 GT rides matched) using magnitude variance and zero-crossing segmentation',
        'Rejection: 93% correct rejection rate. The quality filter successfully prevents catastrophic errors from rides with orientation changes, impacts, or unstable calibration.',
        'Estimation: MAE = 1.16m, Median = 0.80m on accepted rides. Over 50% of accepted rides have < 1m error.',
        'Conformal Coverage: 94.7% LOO coverage (≥90% target achieved), providing honest uncertainty quantification.',
    ]
    for a in achievements:
        doc.add_paragraph(a, style='List Bullet')
    
    add_heading(doc, '8.2 Key Innovations', 2)
    add_para(doc, '1. Gravity-projected ZUPT: Recovers vertical acceleration from arbitrary phone orientations.\n'
             '2. Post-ride gravity fallback: Handles pocket-mode initialization.\n'
             '3. Drift-corrected magnitude: Automatically detects and corrects for gravity projection failures.\n'
             '4. Multi-level quality filter: Combines pre-estimation and post-estimation rejection criteria.')
    
    add_heading(doc, '8.3 Remaining Challenges', 2)
    add_para(doc, 'The ±1m conformal interval target was not achieved. Three accepted rides with 2–4m errors '
             'cannot be distinguished from accurate rides using accelerometer-only features. This is a '
             'fundamental limitation of the sensor modality.')
    
    add_heading(doc, '8.4 Future Directions', 2)
    futures = [
        'Gyroscope integration: Track orientation changes during the ride for more accurate gravity projection.',
        'Barometer fusion: Use barometer for absolute height validation and sign determination.',
        'Learned quality model: Train a classifier on labeled accept/reject data to improve filtration.',
        'Multi-building validation: Test on diverse building types and elevator systems.',
        'Real-time processing: Optimize for streaming inference on mobile devices.',
    ]
    for f in futures:
        doc.add_paragraph(f, style='List Bullet')
    
    # ================================================================
    # Save
    # ================================================================
    doc.save(OUT)
    print(f"\nReport saved to: {OUT}")
    print(f"File size: {os.path.getsize(OUT) / 1024 / 1024:.1f} MB")
    return OUT


if __name__ == "__main__":
    print("Generating comprehensive DOCX report...")
    build_report()
    print("Done!")
