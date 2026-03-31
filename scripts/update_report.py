#!/usr/bin/env python3
"""
Update Research_Report.docx with new sections for:
- Using Pre-Computed Segments
- Custom Dataset Evaluation Guide
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level):
    """Add a heading at the specified level."""
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold=False, italic=False):
    """Add a paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

def add_code_block(doc, code):
    """Add a code-style paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.style = doc.styles['No Spacing'] if 'No Spacing' in [s.name for s in doc.styles] else None
    # Grey background via shading
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F2F2F2',
    })
    pPr.append(shd)
    return p

def add_table_row(table, cells_text, bold=False):
    """Add a row to a table."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(10)
    return row

def find_paragraph_index(doc, heading_text):
    """Find the index of a paragraph by its text content."""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading_text.strip():
            return i
    return -1

def insert_after_heading(doc, after_heading, elements_func):
    """Insert content after a specific heading by rebuilding at the end."""
    # Since python-docx doesn't support easy insertion, we append at the end
    elements_func(doc)

def add_figure(doc, image_path, caption, width=Inches(5.5)):
    """Add a figure with caption."""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"[Figure: {caption} — file not found: {image_path}]", italic=True)

def main():
    report_path = os.path.join("docs", "Research_Report.docx")
    doc = Document(report_path)
    
    # ================================================================
    # Section 7.6: Using Pre-Computed Segments (insert before Limitations)
    # We'll add new sections at the end of chapter 7
    # ================================================================
    
    # Find where "8. Conclusions" starts so we can insert before it
    # Since python-docx doesn't support easy mid-document insertion,
    # we'll add new sections. The simplest robust approach: append to doc.
    # The sections will appear at the end, which we can note.
    
    # Actually, let's find the "7.6 Limitations" heading and add after it,
    # then add the new sections before chapter 8.
    # The cleanest approach: add new top-level sections at the end.
    
    # Add page break before new content
    doc.add_page_break()
    
    # ================================================================
    # NEW SECTION: Custom Evaluation & Pre-Computed Segments
    # ================================================================
    
    add_heading(doc, '9. Custom Evaluation & Pre-Computed Segments Guide', level=1)
    
    add_para(doc, 
        'This chapter documents how to evaluate the pipeline on new datasets '
        'and how to use pre-computed detection/segmentation results. These features '
        'enable users to (1) test the algorithm on their own tagged elevator data, '
        'and (2) bypass the detection stage when segment boundaries are already known.')
    
    # --- 9.1 Pre-Computed Segments ---
    add_heading(doc, '9.1 Using Pre-Computed Segments', level=2)
    
    add_para(doc,
        'When segment boundaries are already known (e.g., from manual annotation or '
        'an external detection algorithm), the detection stage can be skipped entirely. '
        'The pipeline runs only the quality filter and height estimation stages on the '
        'provided segments.')
    
    add_heading(doc, '9.1.1 Python API', level=3)
    
    add_para(doc, 'The process_segments() method accepts user-defined time intervals:')
    
    add_code_block(doc,
        'from src.pipeline import ElevatorHeightPipeline\n\n'
        'pipeline = ElevatorHeightPipeline.load("model/")\n\n'
        'segments = [\n'
        '    {"start_time": 10.5, "end_time": 25.3},\n'
        '    {"start_time": 45.0, "end_time": 62.1},\n'
        ']\n\n'
        '# Skip detection, run quality + estimation only\n'
        'results = pipeline.process_segments(acc_x, acc_y, acc_z, segments, fs=100)\n\n'
        '# Visual output\n'
        'results, fig = pipeline.process_segments_plot(\n'
        '    acc_x, acc_y, acc_z, segments, fs=100, save_path="output.png")')
    
    add_para(doc,
        'The output format is identical to process(), with each result containing '
        'height_estimate, confidence_interval_90, method, accepted status, and quality features.')
    
    add_heading(doc, '9.1.2 Command-Line Interface', level=3)
    
    add_para(doc, 'The --segments flag in run_inference.py accepts a JSON file:')
    
    add_code_block(doc,
        '# Segments JSON file format: [{\"start_time\": 10.5, \"end_time\": 25.3}, ...]\n'
        'python run_inference.py --input data.csv --segments segments.json --verbose')
    
    # --- 9.2 Custom Dataset Evaluation ---
    add_heading(doc, '9.2 Custom Dataset Evaluation', level=2)
    
    add_para(doc,
        'The run_custom_evaluation.py script enables rigorous evaluation of the pipeline '
        'on user-provided tagged datasets. This is designed for test-set evaluation: the '
        'pre-calibrated model is used as-is with no re-calibration on the user\'s data. '
        'This ensures an honest assessment of generalization performance.')
    
    add_heading(doc, '9.2.1 Dataset CSV Format', level=3)
    
    add_para(doc,
        'The user provides a CSV index file where each row describes one elevator ride segment:')
    
    # Create format table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, text in enumerate(['Column', 'Required', 'Type', 'Description']):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    rows_data = [
        ['segment_id', 'Yes', 'int', 'Unique identifier for each ride'],
        ['acc_data_path', 'Yes', 'str', 'Path to accelerometer CSV file'],
        ['start_time', 'Yes', 'float', 'Segment start time (seconds)'],
        ['end_time', 'Yes', 'float', 'Segment end time (seconds)'],
        ['true_height', 'Yes', 'float', 'Ground-truth height difference (m, signed)'],
        ['phone_position', 'No', 'str', 'Phone position (hand, pocket, etc.)'],
        ['fs', 'No', 'float', 'Sampling rate in Hz (default: auto-detect)'],
    ]
    for row_data in rows_data:
        add_table_row(table, row_data)
    
    add_para(doc, '')
    add_para(doc, 'Example CSV:')
    add_code_block(doc,
        'segment_id,acc_data_path,start_time,end_time,true_height,phone_position\n'
        '1,data/recording1.csv,85.9,108.0,18.4,hand\n'
        '2,data/recording1.csv,150.3,172.1,-12.6,hand\n'
        '3,data/recording2.csv,22.0,45.5,6.1,pocket')
    
    add_para(doc,
        'Multiple segments can reference the same accelerometer file. The accelerometer CSV '
        'can use standard column headers (acc_x, acc_y, acc_z) or headerless 4-column format '
        '(time_ms, x, y, z). Timestamps are automatically detected and converted.')
    
    add_heading(doc, '9.2.2 Evaluation Modes', level=3)
    
    add_para(doc, 'Two evaluation modes are available:')
    
    mode_table = doc.add_table(rows=1, cols=3)
    mode_table.style = 'Light Grid Accent 1'
    hdr = mode_table.rows[0].cells
    for i, text in enumerate(['Mode', 'What It Tests', 'When to Use']):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    add_table_row(mode_table, [
        'segments_only',
        'Quality filter + height estimation',
        'You have reliable segment boundaries and want to test estimation accuracy'
    ])
    add_table_row(mode_table, [
        'full',
        'Detection + segmentation + quality + estimation',
        'You want to test the entire pipeline end-to-end'
    ])
    
    add_para(doc, '')
    add_para(doc, 'Command-line usage:')
    add_code_block(doc,
        '# Segments-only (recommended for testing estimation quality)\n'
        'python run_custom_evaluation.py --dataset my_data.csv --mode segments_only -v\n\n'
        '# Full pipeline evaluation\n'
        'python run_custom_evaluation.py --dataset my_data.csv --mode full -v\n\n'
        '# Custom output directory\n'
        'python run_custom_evaluation.py --dataset my_data.csv -o results/ -v')
    
    add_heading(doc, '9.2.3 Output Files', level=3)
    
    add_para(doc,
        'All outputs are saved to the --output-dir directory (default: evaluation_output/):')
    
    out_table = doc.add_table(rows=1, cols=2)
    out_table.style = 'Light Grid Accent 1'
    hdr = out_table.rows[0].cells
    hdr[0].text = 'File'
    hdr[1].text = 'Description'
    for p in hdr[0].paragraphs:
        for run in p.runs:
            run.bold = True
    for p in hdr[1].paragraphs:
        for run in p.runs:
            run.bold = True
    
    output_files = [
        ['results.json', 'Per-ride results: estimated height, error, method, accept/reject, quality features'],
        ['summary.json', 'Aggregate metrics: MAE, median error, acceptance rate, conformal coverage'],
        ['fig_scatter.png', 'True vs Estimated height scatter plot'],
        ['fig_per_ride_errors.png', 'Per-ride absolute error bar chart'],
        ['fig_error_histogram.png', 'Error distribution histogram'],
        ['fig_error_cdf.png', 'Cumulative error distribution'],
        ['fig_conformal_coverage.png', 'Pre-calibrated conformal interval coverage check'],
        ['fig_quality_vs_error.png', 'Quality score vs height error'],
        ['fig_rejection_analysis.png', 'Rejection reasons breakdown'],
        ['fig_error_vs_height.png', 'Error vs ride height magnitude'],
        ['fig_method_breakdown.png', 'Estimation method usage and accuracy'],
        ['fig_individual_rides.png', 'Best/worst ride height curves'],
        ['fig_summary_dashboard.png', 'Summary dashboard with key metrics'],
        ['fig_detection_*.png', 'Detection timeline (full mode only)'],
    ]
    for row_data in output_files:
        add_table_row(out_table, row_data)
    
    add_heading(doc, '9.2.4 Interpreting Results', level=3)
    
    add_para(doc, 'Key metrics to check in the evaluation output:', bold=True)
    
    p = doc.add_paragraph()
    run = p.add_run('1. Conformal Coverage ≥ 90%: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        'The pre-calibrated conformal interval should cover at least 90% of accepted rides '
        'on the test set. If coverage drops significantly below 90%, the pipeline may not '
        'generalize well to the new data distribution.')
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run('2. Accepted MAE: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        'Mean absolute error on accepted rides. Target is ≤1.0m (sub-floor accuracy).')
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run('3. Acceptance Rate: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        'Fraction of rides passing quality filtering. Low acceptance rate indicates '
        'challenging data (e.g., pocket rides, moving phone).')
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run('4. Rejection Quality: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        'What fraction of rejected rides genuinely had high error (>1m). '
        'High values (>70%) indicate the quality filter is correctly identifying unreliable segments.')
    run.font.size = Pt(11)
    
    # --- 9.3 Example: Bar-Ilan Evaluation ---
    add_heading(doc, '9.3 Example: Bar-Ilan Test-Set Evaluation', level=2)
    
    add_para(doc,
        'To validate the evaluation pipeline, we ran it on the Bar-Ilan dataset '
        'in segments_only mode (using ground-truth segment boundaries). This tests '
        'the quality filter and height estimation stages independently of detection quality.')
    
    add_para(doc, 'Commands:')
    add_code_block(doc,
        'python scripts/generate_example_eval_csv.py\n'
        'python run_custom_evaluation.py --dataset datasets/bar_ilan_eval_example.csv \\\n'
        '    --mode segments_only --output-dir evaluation_output/ -v')
    
    # Results table
    add_para(doc, 'Results summary:', bold=True)
    
    res_table = doc.add_table(rows=1, cols=2)
    res_table.style = 'Light Grid Accent 1'
    hdr = res_table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    for p in hdr[0].paragraphs:
        for run in p.runs:
            run.bold = True
    for p in hdr[1].paragraphs:
        for run in p.runs:
            run.bold = True
    
    metrics = [
        ['Total Rides', '33'],
        ['Accepted', '12 (36%)'],
        ['Rejected', '21'],
        ['Accepted MAE', '0.996m'],
        ['Accepted Median Error', '0.579m'],
        ['Max Error (Accepted)', '3.881m'],
        ['Within 1m', '7/12 (58%)'],
        ['Within 2m', '11/12 (92%)'],
        ['Conformal Interval', '±3.62m (pre-calibrated)'],
        ['Conformal Coverage', '91.7% (≥90% target ✓)'],
        ['Rejection Quality', '90% had error >1m'],
    ]
    for row_data in metrics:
        add_table_row(res_table, row_data)
    
    add_para(doc, '')
    
    # Add key figures
    fig_dir = 'evaluation_output'
    
    add_para(doc, 'Selected evaluation figures:', bold=True)
    
    add_figure(doc, os.path.join(fig_dir, 'fig12_summary_dashboard.png'),
               'Figure: Test-set evaluation summary dashboard')
    
    add_figure(doc, os.path.join(fig_dir, 'fig01_scatter.png'),
               'Figure: True vs Estimated height on test set')
    
    add_figure(doc, os.path.join(fig_dir, 'fig02_per_ride_errors.png'),
               'Figure: Per-ride errors on test set (green=accepted, red=rejected)')
    
    add_figure(doc, os.path.join(fig_dir, 'fig05_conformal_coverage.png'),
               'Figure: Pre-calibrated conformal coverage on test set (91.7% ≥ 90% target)')
    
    add_figure(doc, os.path.join(fig_dir, 'fig04_error_cdf.png'),
               'Figure: Cumulative error distribution on test set')
    
    add_figure(doc, os.path.join(fig_dir, 'fig07_rejection_analysis.png'),
               'Figure: Rejection analysis — error distribution and rejection reasons')
    
    add_figure(doc, os.path.join(fig_dir, 'fig10_individual_rides.png'),
               'Figure: Individual ride analysis — best and worst accepted rides')
    
    add_para(doc,
        'The evaluation confirms that the pre-calibrated conformal interval achieves '
        '91.7% coverage on the test set (exceeding the 90% target). The quality filter '
        'correctly rejects 90% of rides that would have had error >1m. Accepted rides '
        'achieve a median error of 0.579m, demonstrating sub-floor accuracy on reliable segments.')
    
    # Save
    doc.save(report_path)
    print(f"Report updated: {report_path}")
    print("Added: Chapter 9 — Custom Evaluation & Pre-Computed Segments Guide")

if __name__ == "__main__":
    main()
