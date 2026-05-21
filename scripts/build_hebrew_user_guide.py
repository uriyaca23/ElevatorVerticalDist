"""Build the Hebrew operator user guide as a Word document.

Output: docs/boutique_pipeline_guide_hebrew.docx

Run from repo root with the project venv active:

    python scripts/build_hebrew_user_guide.py

Hebrew RTL is handled natively by Word: every paragraph that holds
Hebrew text is created via `_p_he()`, which sets bidi=True on the
paragraph and right-aligns it. Image paths are resolved relative to
REPO so the script is location-stable.

Screenshots that don't exist yet are written as inline Hebrew
placeholders (`[צילום מסך: ... — להוסיף]`) so the user can capture and
re-run the script later.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
FIG = REPO / "docs" / "latex" / "figures"
BOUTIQUE = FIG / "boutique"
PANELS = BOUTIQUE / "panels"
CROPS = BOUTIQUE / "crops"
PAPER_FIG = REPO / "paper_phd" / "figures"
OUTPUT = REPO / "docs" / "boutique_pipeline_guide_hebrew.docx"


HEBREW_FONT = "David"
LATIN_FONT = "Calibri"


# ---------------------------------------------------------------------------
# Low-level RTL helpers
# ---------------------------------------------------------------------------

def _set_paragraph_bidi(paragraph) -> None:
    """Mark a paragraph as right-to-left so Word reorders runs correctly."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)


def _set_run_rtl(run) -> None:
    """Mark a run as Hebrew/RTL so Word picks the Hebrew font slot."""
    rPr = run._r.get_or_add_rPr()
    rtl = rPr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)
    # Use the cs (complex script) font slot for Hebrew. python-docx's
    # Run.font.name only sets the ASCII slot, which Word ignores for RTL.
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), HEBREW_FONT)
    rFonts.set(qn("w:ascii"), HEBREW_FONT)
    rFonts.set(qn("w:hAnsi"), HEBREW_FONT)


def _p_he(doc, text: str, *, bold: bool = False, size: int = 12,
          align=WD_ALIGN_PARAGRAPH.RIGHT, color: RGBColor | None = None):
    """Add a single-run Hebrew (RTL) paragraph."""
    p = doc.add_paragraph()
    p.alignment = align
    _set_paragraph_bidi(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    _set_run_rtl(run)
    return p


def _p_he_multi(doc, parts: list[tuple[str, dict]],
                align=WD_ALIGN_PARAGRAPH.RIGHT):
    """Add an RTL paragraph composed of multiple runs.

    Each part is (text, kwargs) where kwargs accept: bold, italic,
    size, color, latin (bool — if True, use Calibri instead of David).
    """
    p = doc.add_paragraph()
    p.alignment = align
    _set_paragraph_bidi(p)
    for text, kw in parts:
        run = p.add_run(text)
        run.bold = bool(kw.get("bold"))
        run.italic = bool(kw.get("italic"))
        if "size" in kw:
            run.font.size = Pt(kw["size"])
        if "color" in kw:
            run.font.color.rgb = kw["color"]
        if kw.get("latin"):
            run.font.name = LATIN_FONT
        else:
            _set_run_rtl(run)
    return p


def _heading_he(doc, text: str, level: int = 1):
    """Hebrew heading. Uses Word's built-in Heading styles for TOC."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_bidi(h)
    for run in h.runs:
        _set_run_rtl(run)
    return h


def _page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _picture(doc, rel_path: Path, width_cm: float = 15.5, caption: str | None = None,
             missing_note_he: str | None = None):
    """Embed a picture if it exists; otherwise write a Hebrew placeholder.

    rel_path is relative to REPO root.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if rel_path.exists():
        run = p.add_run()
        run.add_picture(str(rel_path), width=Cm(width_cm))
    else:
        note = missing_note_he or f"[חסר: {rel_path.name}]"
        _p_he(doc, note, size=11, color=RGBColor(0xB0, 0x00, 0x00),
              align=WD_ALIGN_PARAGRAPH.CENTER)
    if caption:
        _p_he(doc, caption, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)


def _placeholder_screenshot(doc, what_he: str):
    """A clearly marked Hebrew placeholder for screenshots to add later."""
    note = f"[צילום מסך להוספה: {what_he}]"
    _p_he(doc, note, size=11, color=RGBColor(0xB0, 0x60, 0x00),
          align=WD_ALIGN_PARAGRAPH.CENTER)


# ---------------------------------------------------------------------------
# Sections
# ---------------------------------------------------------------------------

def _setup_section(doc):
    """Default section: A4, 2 cm margins, bidi=True on sectPr."""
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    sectPr = section._sectPr
    bidi = sectPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        sectPr.append(bidi)


def _setup_styles(doc):
    """Make the Normal style Hebrew-friendly."""
    normal = doc.styles["Normal"]
    normal.font.name = HEBREW_FONT
    normal.font.size = Pt(12)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), HEBREW_FONT)


def _write_title_page(doc):
    for _ in range(4):
        doc.add_paragraph()

    _p_he(doc, "מדריך למשתמש",
          bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p_he(doc, "מערכת לאומדן תזוזה אנכית של מעלית באמצעות סמארטפון",
          bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Boutique Pipeline — User Guide")
    r.font.name = LATIN_FONT
    r.font.size = Pt(14)
    r.italic = True

    for _ in range(3):
        doc.add_paragraph()

    _p_he(doc, "מחברים:", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Eyal Yakir   •   Uriya Cohen-Eliya")
    r.font.name = LATIN_FONT
    r.font.size = Pt(13)

    doc.add_paragraph()
    _p_he(doc, "חוג / מחלקה: [להשלים]", size=12,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          color=RGBColor(0x60, 0x60, 0x60))
    _p_he(doc, "מוסד: [Reichman University / Tel Aviv University — להשלים]",
          size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
          color=RGBColor(0x60, 0x60, 0x60))

    for _ in range(2):
        doc.add_paragraph()

    _p_he(doc, f"תאריך: {datetime.today().strftime('%d/%m/%Y')}",
          size=12, align=WD_ALIGN_PARAGRAPH.CENTER)


def _write_toc(doc):
    _heading_he(doc, "תוכן עניינים", level=1)

    # Live TOC field (Word will populate on first "Update Field")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_bidi(p)
    r = p.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    r._r.append(fldChar_begin)
    r._r.append(instrText)
    r._r.append(fldChar_sep)
    r._r.append(fldChar_end)

    doc.add_paragraph()
    _p_he(doc, "(ב-Word: לחיצה ימנית על הטבלה מעלה אפשרות 'Update Field' שתעדכן את מספרי העמודים.)",
          size=10, color=RGBColor(0x60, 0x60, 0x60))

    # Static fallback list
    doc.add_paragraph()
    _p_he(doc, "סקירת תוכן:", bold=True, size=12)
    for item in [
        "1. סקירה כללית של התהליך",
        "2. שלב טעינת הנתונים",
        "3. שלב הסגמנטציה (זיהוי נסיעות)",
        "4. שלב חיזוי הגובה",
        "5. שלב הורדת הדו\"ח",
    ]:
        _p_he(doc, item, size=12)


def _write_workflow_overview(doc):
    _heading_he(doc, "סקירה כללית — איך עובד התהליך", level=1)

    _p_he(doc,
          "המערכת מקבלת הקלטה גולמית של מד-תאוצה מהטלפון של הנוסע, "
          "מזהה את קטעי הזמן בהם הנוסע היה בתוך מעלית נעה, ומעריכה לכל "
          "נסיעה את התזוזה האנכית (Δh) במטרים יחד עם רווח-בר-סמך של 90%.")
    _p_he(doc,
          "התהליך מורכב מארבעה שלבים. בשניים מהם המשתמש פעיל ובוחן את "
          "תוצאות המערכת; בשניים האחרים אין צורך בהתערבות.")

    # Workflow table (4 columns, 1 row of step boxes + 1 row of arrows)
    table = doc.add_table(rows=2, cols=7)
    # Step / arrow / step / arrow / step / arrow / step
    cells = table.rows[0].cells
    step_names = [
        ("1. טעינת נתונים", False),
        ("→", None),
        ("2. סגמנטציה", True),
        ("→", None),
        ("3. חיזוי גובה", True),
        ("→", None),
        ("4. הורדת דו\"ח", False),
    ]
    for cell, (text, interactive) in zip(cells, step_names):
        cell.text = ""  # clear
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        if interactive is not None:
            _set_paragraph_bidi(p)
            _set_run_rtl(run)
            if interactive:
                run.font.color.rgb = RGBColor(0xC0, 0x60, 0x00)
        else:
            run.font.name = LATIN_FONT
            run.font.size = Pt(16)

    # Row of role labels under each step
    role_row = table.rows[1].cells
    roles = [
        "לא מפוקח",
        "",
        "סקירה ועריכה",
        "",
        "אימות התאמת מודל",
        "",
        "לא מפוקח",
    ]
    for cell, text in zip(role_row, roles):
        cell.text = ""
        if not text:
            continue
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_bidi(p)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        _set_run_rtl(run)

    doc.add_paragraph()
    _p_he(doc, "תפקיד כל שלב:", bold=True, size=13)
    _p_he_multi(doc, [
        ("שלב 1 — ", {"bold": True}),
        ("טעינת הנתונים מהטלפון או מקובץ Excel. שלב טכני, ללא בדיקה ידנית.", {}),
    ])
    _p_he_multi(doc, [
        ("שלב 2 — ", {"bold": True}),
        ("האלגוריתם מזהה אוטומטית את קטעי הנסיעות. ", {}),
        ("כאן עליך לבדוק את התוצאות ולערוך אם צריך", {"bold": True, "color": RGBColor(0xC0, 0x60, 0x00)}),
        (": למחוק זיהויים שגויים, לתקן גבולות, ולהוסיף נסיעות שהמערכת פספסה.", {}),
    ])
    _p_he_multi(doc, [
        ("שלב 3 — ", {"bold": True}),
        ("שני אלגוריתמים מחשבים את הגובה לכל נסיעה. ", {}),
        ("כאן עליך לוודא שהמודל המותאם נראה טוב על האות", {"bold": True, "color": RGBColor(0xC0, 0x60, 0x00)}),
        (", ולתקן ידנית פרמטרים במקרה הצורך.", {}),
    ])
    _p_he_multi(doc, [
        ("שלב 4 — ", {"bold": True}),
        ("הורדת קובץ PDF סופי. שלב טכני, ללא בדיקה ידנית.", {}),
    ])


def _write_step1_data(doc):
    _heading_he(doc, "שלב 1 — טעינת הנתונים", level=1)

    _p_he(doc, "ישנן שתי דרכים להזין נתונים למערכת:", bold=True, size=13)

    _heading_he(doc, "א. טעינה מטלפון לפי ID + טווח זמן", level=2)
    _p_he(doc,
          "אם הטלפון נמצא במאגר הפנימי, ניתן לבחור את ה-ID שלו ולציין טווח "
          "זמן (התחלה וסיום). המערכת תוריד אוטומטית את נתוני מד-התאוצה "
          "לטווח שביקשת.")
    _p_he_multi(doc, [
        ("שים לב — ", {"bold": True, "color": RGBColor(0xB0, 0x00, 0x00)}),
        ("שאילתת 15 דקות אורכת כ-5 דקות עיבוד. ", {"bold": True}),
        ("השתדל להגדיר טווחי זמן קצרים ככל הניתן (לדוגמה: רק חצי השעה "
         "שבה התקיימו הנסיעות), אחרת תיאלץ להמתין זמן רב.", {}),
    ])

    _heading_he(doc, "ב. העלאת קובץ CSV / Excel", level=2)
    _p_he(doc,
          "כל מי שאינו במאגר הפנימי יכול להעלות קובץ ידנית. הקובץ חייב "
          "להכיל שתי עמודות עם כותרת בשורה הראשונה:")
    _p_he_multi(doc, [
        ("• ", {}),
        ("time", {"latin": True, "bold": True}),
        (" — זמן בשניות או במילישניות (זיהוי אוטומטי).", {}),
    ])
    _p_he_multi(doc, [
        ("• ", {}),
        ("vertical_acceleration", {"latin": True, "bold": True}),
        (" — מד-תאוצה אנכי ביחידות של m/s². ", {}),
    ])

    doc.add_paragraph()
    _p_he(doc, "מסך הבחירה במערכת:", bold=True, size=12)
    _picture(doc, BOUTIQUE / "ui_step2_picker.png", width_cm=13,
             caption="בחירה בין שני מקורות הנתונים.")
    _picture(doc, BOUTIQUE / "ui_step2_upload_form.png", width_cm=13,
             caption="טופס העלאת קובץ.")

    doc.add_paragraph()
    _p_he_multi(doc, [
        ("שלב זה אינו דורש התערבות. ", {"bold": True, "color": RGBColor(0x00, 0x70, 0x00)}),
        ("לאחר טעינת הנתונים, לוחצים 'Next' ועוברים לשלב הסגמנטציה.", {}),
    ])


def _write_step2_segmentation(doc):
    _heading_he(doc, "שלב 2 — סגמנטציה (זיהוי נסיעות)", level=1)

    # ---- 2a: theory
    _heading_he(doc, "רקע — מה אנחנו מנסים לזהות?", level=2)
    _p_he(doc,
          "מעלית מודרנית אינה זזה באופן שרירותי. תא המעלית כפוף לשלוש מגבלות "
          "פיזיקליות שמטיל היצרן:")
    _p_he_multi(doc, [("• ", {}),
                      ("הגבלת ", {}),
                      ("jerk", {"latin": True, "italic": True}),
                      (" — קצב השינוי של התאוצה (חוויית רכיבה חלקה).", {})])
    _p_he_multi(doc, [("• ", {}),
                      ("הגבלת תאוצה מקסימלית ", {}),
                      ("a_max", {"latin": True, "italic": True}),
                      (" — היכולת של המנוע להאיץ.", {})])
    _p_he_multi(doc, [("• ", {}),
                      ("הגבלת מהירות מקסימלית ", {}),
                      ("v_max", {"latin": True, "italic": True}),
                      (" — מהירות שיוט מותרת.", {})])

    _p_he(doc,
          "התוצאה היא שהמערכת יודעת מראש שצורת אות מד-התאוצה בכל נסיעה "
          "שייכת למשפחה צרה של צורות (פולסים) — והדבר מאפשר לחפש אותם "
          "באות הגולמי.")

    # Boxed θ identity
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "6")
        b.set(qn("w:color"), "808080")
        pBdr.append(b)
    pPr.append(pBdr)
    r = p.add_run("θ = ( j_max ,  a_max ,  v_max ,  H )")
    r.font.name = LATIN_FONT
    r.font.size = Pt(14)
    r.bold = True

    _p_he(doc,
          "כאשר שלוש המגבלות הראשונות (j_max, a_max, v_max) הן תכונות "
          "קבועות של המעלית, ורק הגובה H (המרחק האנכי שעוברת המעלית) "
          "משתנה מנסיעה לנסיעה.")

    _p_he(doc, "שלושה משטרי תנועה אפשריים:", bold=True, size=12)
    _p_he_multi(doc, [
        ("• ", {}),
        ("נסיעה ארוכה — 7 קטעים: ", {"bold": True}),
        ("המעלית מספיקה להגיע למהירות שיוט מלאה. שני 'אונות' (lobes) של "
         "תאוצה במרכזים שטוחים, מופרדות על-ידי שלב שיוט שטוח באמצע.", {}),
    ])
    _p_he_multi(doc, [
        ("• ", {}),
        ("נסיעה בינונית — 5 קטעים: ", {"bold": True}),
        ("התאוצה מגיעה למקסימום אבל המהירות לא. שתי האונות נוגעות זו בזו, "
         "ללא שלב שיוט באמצע.", {}),
    ])
    _p_he_multi(doc, [
        ("• ", {}),
        ("נסיעה קצרה — 3 קטעים: ", {"bold": True}),
        ("אפילו התאוצה לא מגיעה למקסימום. שתי האונות הן משולשים שנוגעים "
         "באמצע.", {}),
    ])

    _picture(doc, FIG / "triangle_vs_trapezoid.png", width_cm=14,
             caption="שלושת משטרי הנסיעה: שמאל — קצרה (משולשים), אמצע — בינונית, ימין — ארוכה (טרפזים).")

    _p_he(doc, "כך נראים פולסים אמיתיים בהקלטות אמת:", size=12)
    _picture(doc, FIG / "elevator_kinematics_all_rides.png", width_cm=15.5,
             caption="כל הנסיעות מהקלטה אחת — בכל אחת מהן צורת 'שתי אונות' חוזרת. רוחב המרווח בין האונות תלוי בגובה H.")

    # ---- 2b: what the step shows
    _heading_he(doc, "מה מציג השלב במערכת?", level=2)
    _p_he(doc,
          "המערכת מריצה את הגלאי באופן אוטומטי ומציגה את האות המלא יחד עם "
          "הקטעים שזיהתה. כל קטע שזוהה מסומן ברצועה צבעונית: ")
    _p_he_multi(doc, [
        ("• כחול = ", {}),
        ("עלייה (up)", {"bold": True, "color": RGBColor(0x00, 0x60, 0xC0)}),
    ])
    _p_he_multi(doc, [
        ("• מגנטה = ", {}),
        ("ירידה (down)", {"bold": True, "color": RGBColor(0xB0, 0x30, 0x80)}),
    ])
    _p_he_multi(doc, [
        ("• כתום = ", {}),
        ("הקטע הנבחר כרגע (לצורך עריכה / עיון)", {"bold": True, "color": RGBColor(0xD0, 0x80, 0x00)}),
    ])
    _picture(doc, PANELS / "clean_step3_overview__signal.png", width_cm=15.5,
             caption="האות המלא עם הקטעים שזוהו אוטומטית — דוגמה להקלטה נקייה עם 12 נסיעות ברורות.")

    # ---- 2c: Task A — verify good
    _heading_he(doc, "המשימה הראשונה — לוודא שכל קטע 'טוב'", level=2)
    _p_he(doc,
          "לכל קטע, המערכת מציגה ארבעה פאנלים אבחנתיים. הפאנל החשוב ביותר "
          "הוא ההתאמה של תבנית הטרפז (red+purple) על האות הגולמי. קטע "
          "תקין נראה כך:")

    _p_he(doc, "שלוש דוגמאות לקטעים תקינים:", bold=True, size=13,
          color=RGBColor(0x00, 0x70, 0x00))

    _picture(doc, PANELS / "clean_step3_overview__trap.png", width_cm=15,
             caption="✓ התאמת טרפז תקינה: שתי האונות יושבות על מעטפת האות; הפלאטו ≈ ±0.5 m/s² (טווח אופייני).")
    _picture(doc, PANELS / "clean_step3_overview__corr.png", width_cm=15,
             caption="✓ פאנל קורלציה תקין: שתי נקודות ירוקות (אחת חיובית, אחת שלילית) המייצגות זוג אונות מאומת.")
    _picture(doc, PANELS / "clean_step3_overview__heatmaps.png", width_cm=15,
             caption="✓ Heatmap תקין: ה-X (אופטימום) נמצא במרכז הרשת, האזור הבהיר קומפקטי. הקטע מוגדר היטב.")

    _p_he(doc, "שלוש דוגמאות לקטעים בעייתיים:", bold=True, size=13,
          color=RGBColor(0xB0, 0x00, 0x00))

    _picture(doc, PANELS / "fp_step3_seg0_FP__trap.png", width_cm=15,
             caption="✗ False positive — אמפליטודה ענקית (50 m/s² במקום 1 m/s²). זוהי הקשה/מכה על הטלפון, לא נסיעת מעלית. למחוק.")
    _picture(doc, PANELS / "damped_step3_seg0_full__trap.png", width_cm=15,
             caption="✗ הקלטה דעוכה — האונות בקושי מעל הרעש (פלאטו ≈ ±0.4 m/s²). זו אכן נסיעה אמיתית, אך אמפליטודה חלשה כי הטלפון בכיס/תיק. רווח-הסמך יתרחב אוטומטית.")
    _picture(doc, PANELS / "haari_step3_seg20_full__trap.png", width_cm=15,
             caption="✗ קטע ארוך מדי (איחוד) — הקטע משתרע על 22 שניות וכולל שתי נסיעות נפרדות עם שקט ארוך באמצע. לחתוך לשני קטעים נפרדים.")

    # ---- 2d: Task A continued — edit/delete
    _heading_he(doc, "תיקון — מחיקה ועריכה", level=2)
    _p_he(doc,
          "בסיידבר השמאלי של המסך יש ארבעה כלי-עריכה עיקריים: ")
    _p_he_multi(doc, [
        ("• ", {}),
        ("Edit selected", {"latin": True, "bold": True}),
        (" — חלון קופץ עם זמן התחלה/סיום וסוג (up/down).", {}),
    ])
    _p_he_multi(doc, [
        ("• ", {}),
        ("Delete", {"latin": True, "bold": True}),
        (" — אייקון פח אשפה למחיקת הקטע הנבחר.", {}),
    ])
    _p_he_multi(doc, [
        ("• ", {}),
        ("+ Add segment", {"latin": True, "bold": True}),
        (" — הוספת קטע חדש של 5 שניות (סוג 'up' כברירת-מחדל) שאפשר אז לערוך.", {}),
    ])
    _p_he_multi(doc, [
        ("• ", {}),
        ("Reset to detector output", {"latin": True, "bold": True}),
        (" — איפוס כל העריכות הידניות והחזרה לפלט המקורי של הגלאי.", {}),
    ])

    _placeholder_screenshot(doc, "לחיצה על אייקון הפח למחיקת FP שנבחר")
    _placeholder_screenshot(doc, "חלון 'Edit selected' עם שדות start / end / type")

    _heading_he(doc, "דפוסי תיקון נפוצים", level=3)
    _p_he_multi(doc, [
        ("False positive (FP) — ", {"bold": True}),
        ("רצועה מוצללת מעל אזור שלא באמת מכיל נסיעה (מכה על הטלפון, "
         "טריקת דלת, רעש HVAC). למחוק עם אייקון הפח.", {}),
    ])
    _p_he_multi(doc, [
        ("שתי נסיעות שאוחדו — ", {"bold": True}),
        ("רצועה אחת ארוכה מאוד שמכסה שתי נסיעות אמיתיות עם שקט באמצע. "
         "פותחים Edit selected, חותכים את הסוף לאחר הזוג הראשון של אונות, "
         "ואז + Add segment עבור הנסיעה השנייה.", {}),
    ])
    _p_he_multi(doc, [
        ("הסטה של 1–2 שניות בקצוות — ", {"bold": True}),
        ("הרצועה הכתומה יושבת על הנסיעה הנכונה אבל ההתחלה/סיום קצת מוקדם "
         "או מאוחר מדי. ב-Edit selected מזיזים את הגבול במספר שניות; "
         "ההתאמה מתעדכנת מיידית במסך.", {}),
    ])

    # ---- 2e: Task B — find missed pulses
    _heading_he(doc, "המשימה השנייה — איתור פולסים שהמערכת פספסה", level=2)
    _p_he(doc,
          "הגלאי מכויל להיות שמרני: עדיף לפספס נסיעה מאשר לזהות נסיעה "
          "שגויה. לכן יש לוודא ידנית — על-ידי סריקה בעין של גרף האות "
          "השלם — שאין צורות 'שתי אונות' שלא מוצללות.")
    _picture(doc, PANELS / "damped_step3_overview__signal.png", width_cm=15.5,
             caption="דוגמה להקלטה דעוכה: הגלאי זיהה 10 קטעים בלבד, אך בעין רואים מספר 'אונות זוגיות' נוספות שלא הוצללו — אלו פולסים שפוספסו.")
    _p_he_multi(doc, [
        ("אם זיהית פולס שלא הוצלל: ", {"bold": True}),
        ("לחץ ", {}),
        ("+ Add segment", {"latin": True, "bold": True}),
        (", ואז ", {}),
        ("Edit selected", {"latin": True, "bold": True}),
        (" כדי לקבוע את הגבולות (כמה שניות לפני הפולס הראשון ועד אחרי "
         "הפולס השני) וסוג הנסיעה (up/down).", {}),
    ])
    _placeholder_screenshot(doc, "לחיצה על '+ Add segment' בסיידבר השמאלי")

    _p_he_multi(doc, [
        ("כלל אצבע: ", {"bold": True, "color": RGBColor(0x60, 0x60, 0x60)}),
        ("אם אתה לא ", {}),
        ("רואה בבירור", {"italic": True}),
        (" נסיעה שפוספסה — עדיף לא לערוך. עריכת-יתר על קטעי גבול היא "
         "הטעות הנפוצה ביותר.", {}),
    ])


def _write_step3_prediction(doc):
    _heading_he(doc, "שלב 3 — חיזוי גובה", level=1)

    _p_he(doc,
          "המערכת מריצה שני אלגוריתמים על רשימת הקטעים הסופית: "
          "Trapezoid pulse-pair (כחול) ו-ZUPT (ירוק). שניהם מחזירים אומדן "
          "של Δh לכל קטע. אם שני האלגוריתמים מסכימים — הביטחון בתוצאה גבוה.")

    _picture(doc, CROPS / "clean_step4_metrics.png", width_cm=15.5,
             caption="גרף ההשוואה — לכל קטע מוצגות שתי עמודות (כחול וירוק). הסכמה בגובה דומה ⇐ ביטחון גבוה.")

    _heading_he(doc, "בדיקה לכל קטע — האם ההתאמה נראית טוב?", level=2)
    _p_he(doc,
          "עבור כל קטע, המערכת מציגה את האות הגולמי יחד עם תבנית הטרפז "
          "המותאמת (red+purple). זוהי הבדיקה הקריטית: ההתאמה צריכה לשבת "
          "כמו 'סרגל-עקול' על מעטפת האות.")

    _picture(doc, PANELS / "clean_step3_overview__trap.png", width_cm=15,
             caption="התאמה טובה: שתי האונות יושבות על האות בצורה הדוקה. אין צורך לערוך.")

    _heading_he(doc, "מה לעשות אם אחת מהאונות 'פגומה'?", level=2)
    _p_he(doc,
          "לעיתים אחת האונות נחתכה (קצה הקטע לא תפס אותה במלואה) או "
          "הופרעה (רעש, מכה). במקרה כזה אפשר להשפיע ידנית על שלושת "
          "הפרמטרים של הטרפז:")
    _p_he_multi(doc, [("• ", {}),
                      ("W", {"latin": True, "bold": True}),
                      (" — חצי-רוחב האונה (משך הזמן שלוקח לעלות ולרדת).", {})])
    _p_he_multi(doc, [("• ", {}),
                      ("f", {"latin": True, "bold": True}),
                      (" — חלק הזמן בו הטרפז 'שטוח' למעלה (יחסית לרוחב הכולל).", {})])
    _p_he_multi(doc, [("• ", {}),
                      ("A", {"latin": True, "bold": True}),
                      (" — אמפליטודת הטרפז (גובה הפלאטו).", {})])
    _placeholder_screenshot(doc, "סליידרים של W / f / A לקטע ספציפי בשלב 3")

    _p_he(doc,
          "לאחר כיוון ידני של הפרמטרים, ההתאמה מתעדכנת מיידית בגרף — "
          "אם הטרפז יושב טוב יותר על האות, התוצאה אמינה.")

    _heading_he(doc, "טבלה מפורטת לכל קטע", level=2)
    _picture(doc, CROPS / "clean_step4_top.png", width_cm=15.5,
             caption="לכל קטע מוצגים: Δh חתום, רווח-בר-סמך 90%, ציון איכות, ודגל accept/reject. אם accept=no, ראה את העמודה 'reject_reason'.")

    _p_he(doc,
          "כאשר אתה מרוצה מההתאמה לכל הקטעים, לחץ 'Generate report' "
          "ועבור לשלב הסופי.")


def _write_step4_download(doc):
    _heading_he(doc, "שלב 4 — הורדת הדו\"ח", level=1)

    _p_he_multi(doc, [
        ("שלב זה אינו דורש שום התערבות. ",
         {"bold": True, "color": RGBColor(0x00, 0x70, 0x00)}),
        ("המערכת מציגה תצוגה מקדימה של ה-PDF, ולחיצה על 'Download "
         "Summary PDF' שומרת את הקובץ אצלך.", {}),
    ])

    _picture(doc, CROPS / "clean_step5_top.png", width_cm=15,
             caption="לחיצה על 'Download Summary PDF' שומרת את הדו\"ח.")

    _heading_he(doc, "מה כולל הדו\"ח?", level=2)
    _p_he(doc, "הדו\"ח הוא קובץ PDF רב-עמודים בעברית:")
    _p_he_multi(doc, [
        ("• ", {}),
        ("עמוד שער — ", {"bold": True}),
        ("פרטי הקלטה, סיכום, גרף סקירה, וטבלה לכל הנסיעות.", {}),
    ])
    _p_he_multi(doc, [
        ("• ", {}),
        ("עמוד לכל נסיעה — ", {"bold": True}),
        ("Δh, רווח-סמך, איכות, דגל accept, גרף ההתאמה, וחלון קורלציה ±30 שניות.", {}),
    ])
    _p_he_multi(doc, [
        ("• ", {}),
        ("עמוד מילון — ", {"bold": True}),
        ("הסבר בעברית של כל המדדים, כך שמי שיקבל את הדו\"ח יוכל לקרוא "
         "אותו ללא הסתמכות על המדריך הזה.", {}),
    ])

    doc.add_paragraph()
    _p_he(doc, "שם הקובץ:", bold=True, size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_bidi(p)
    r = p.add_run("boutique_report_YYYYMMDDTHHMMSSZ.pdf")
    r.font.name = LATIN_FONT
    r.font.size = Pt(12)


def _write_footer(doc):
    """Add page numbers + doc title to the footer of every page."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_bidi(p)
    # Page number field
    r = p.add_run("עמוד ")
    _set_run_rtl(r)
    r2 = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r2._r.append(fld_begin)
    r2._r.append(instr)
    r2._r.append(fld_end)
    r3 = p.add_run("   |   מדריך למשתמש — Boutique Pipeline")
    _set_run_rtl(r3)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()
    _setup_section(doc)
    _setup_styles(doc)

    _write_title_page(doc)
    _page_break(doc)

    _write_toc(doc)
    _page_break(doc)

    _write_workflow_overview(doc)
    _page_break(doc)

    _write_step1_data(doc)
    _page_break(doc)

    _write_step2_segmentation(doc)
    _page_break(doc)

    _write_step3_prediction(doc)
    _page_break(doc)

    _write_step4_download(doc)

    _write_footer(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"wrote {OUTPUT}  ({OUTPUT.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    main()
